"""
Generate a professional Word document for the UNIDO GCP Infrastructure Proposal.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Colors
PRIMARY = RGBColor(0x1a, 0x73, 0xe8)
PRIMARY_DARK = RGBColor(0x0d, 0x47, 0xa1)
ACCENT_GREEN = RGBColor(0x34, 0xa8, 0x53)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_PRIMARY = RGBColor(0x20, 0x21, 0x24)
TEXT_SECONDARY = RGBColor(0x5f, 0x63, 0x68)
HEADER_BG = "1a73e8"
ROW_ALT = "f1f3f4"
LIGHT_BG = "e8f0fe"
GREEN_BG = "e6f4ea"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "dadce0")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Bold the last row if it's a total
            is_total = r_idx == len(rows) - 1 and ("total" in str(cell_text).lower() or "total" in str(row[0]).lower())
            run = p.add_run(str(cell_text))
            run.bold = is_total
            run.font.size = Pt(8.5)
            run.font.name = "Calibri"
            run.font.color.rgb = TEXT_PRIMARY
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ROW_ALT)

    # Set column widths
    if col_widths:
        for r_idx, row in enumerate(table.rows):
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = Cm(width)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()
    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY_DARK
        run.font.name = "Calibri"
    return h


def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_b.font.name = "Calibri"
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return p


def add_highlight_box(doc, text, bg_color=LIGHT_BG):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    for line in text.split("\n"):
        p = cell.paragraphs[0] if cell.paragraphs[0].text == "" and text.split("\n").index(line) == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = TEXT_PRIMARY
    set_cell_shading(cell, bg_color)
    set_cell_border(cell,
        top={"color": "1a73e8", "sz": "4"},
        bottom={"color": "1a73e8", "sz": "4"},
        left={"color": "1a73e8", "sz": "12"},
        right={"color": "1a73e8", "sz": "4"})
    doc.add_paragraph()


def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.size = Pt(8)
    run.font.name = "Consolas"
    run.font.color.rgb = TEXT_PRIMARY
    set_cell_shading(cell, "f8f9fa")
    set_cell_border(cell,
        top={"color": "dadce0"}, bottom={"color": "dadce0"},
        left={"color": "dadce0"}, right={"color": "dadce0"})
    doc.add_paragraph()


def add_thin_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="dadce0"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = TEXT_PRIMARY

    # ════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    # Blue line
    p_line = doc.add_paragraph()
    run = p_line.add_run("━" * 80)
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(6)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("UNIDO AI Chatbot")
    run.font.size = Pt(32)
    run.font.color.rgb = PRIMARY_DARK
    run.bold = True
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Google Cloud Platform\nInfrastructure Proposal")
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY
    run.bold = True
    run.font.name = "Calibri"

    doc.add_paragraph()

    # Green accent line
    p_acc = doc.add_paragraph()
    run = p_acc.add_run("━" * 15)
    run.font.color.rgb = ACCENT_GREEN
    run.font.size = Pt(8)

    doc.add_paragraph()

    for line in [
        "Document Version: 1.0",
        "Date: April 29, 2026",
        "Prepared For: UNIDO Technical Team",
        "Classification: Confidential",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT_SECONDARY
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_paragraph()

    # Key metrics table on cover
    metrics_table = doc.add_table(rows=2, cols=3)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in metrics_table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BG)
    for cell in metrics_table.rows[1].cells:
        set_cell_shading(cell, LIGHT_BG)

    metrics = [
        ("Estimated Monthly Cost", "$210 – $325/mo", "(Option A: Gemini Flash)"),
        ("Migration Timeline", "5 – 8 Days", "(Working days)"),
        ("Code Change Effort", "~40% of Codebase", "(60% unchanged)"),
    ]
    for i, (label, value, note) in enumerate(metrics):
        cell = metrics_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_SECONDARY
        run.font.name = "Calibri"
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(value)
        run2.font.size = Pt(16)
        run2.font.color.rgb = PRIMARY_DARK
        run2.bold = True
        run2.font.name = "Calibri"

        cell2 = metrics_table.rows[1].cells[i]
        cell2.text = ""
        p3 = cell2.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(note)
        run3.font.size = Pt(8)
        run3.font.color.rgb = TEXT_SECONDARY
        run3.italic = True
        run3.font.name = "Calibri"

    # Border for metrics
    tbl = metrics_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="1a73e8"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="1a73e8"/>'
        '  <w:left w:val="single" w:sz="6" w:space="0" w:color="1a73e8"/>'
        '  <w:right w:val="single" w:sz="6" w:space="0" w:color="1a73e8"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="dadce0"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    add_thin_line(doc)
    toc_items = [
        "1.  Executive Summary",
        "2.  Current Architecture Overview",
        "3.  Proposed GCP Architecture",
        "4.  GCP Service Mapping — Detailed",
        "5.  Infrastructure Cost Breakdown",
        "6.  Required Code Changes",
        "7.  Migration Timeline",
        "8.  Architecture Diagrams",
        "9.  Recommendations",
        "Appendix A: Environment Variables",
        "Appendix B: GCP CLI Setup Commands",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"●   {item}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = TEXT_PRIMARY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)
    add_thin_line(doc)
    add_body(doc,
        "This document outlines the proposed Google Cloud Platform (GCP) infrastructure for the UNIDO AI Chatbot "
        "application. The project currently uses Azure OpenAI, Elasticsearch, MongoDB, and Azure Blob Storage. "
        "This proposal maps each service to its GCP equivalent, provides detailed cost estimates (one-time and "
        "monthly), and outlines the code changes and timeline required for migration."
    )

    add_highlight_box(doc,
        "Key Highlights:\n"
        "•  Estimated monthly cost: $210 – $325/month (Gemini Flash option)\n"
        "•  One-time migration cost: $150 – $250 (setup and data migration)\n"
        "•  Development effort for code changes: 5 – 8 working days\n"
        "•  Zero downtime migration possible with phased approach\n"
        "•  60% of existing codebase requires no changes"
    )

    # ════════════════════════════════════════════════════════════════
    # 2. CURRENT ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Current Architecture Overview", level=1)
    add_thin_line(doc)
    add_body(doc, "The following table summarizes all current external services and their roles in the application:")

    add_styled_table(doc,
        ["Component", "Current Service", "Purpose"],
        [
            ["AI Chat Model", "Azure OpenAI (GPT-4o-mini)", "Generate conversational responses"],
            ["Embeddings", "Azure OpenAI (text-embedding-3-small)", "Generate 1536D vectors for RAG"],
            ["Vector Search", "Elasticsearch (managed/self-hosted)", "kNN similarity search on embeddings"],
            ["Database", "MongoDB (Atlas or self-hosted)", "Store sessions, logs, admin settings"],
            ["Object Storage", "Azure Blob Storage", "Backup scraped data (JSON files)"],
            ["Web Scraping", "Puppeteer (headless Chrome)", "Daily scraping of UNIDO careers site"],
            ["Frontend", "Static React 19 SPA (Vite)", "User-facing chatbot + admin panel"],
            ["Backend", "Node.js + Express 5", "REST API + WebSocket server"],
            ["Real-time", "Socket.io", "Live admin dashboard updates"],
            ["Scheduling", "node-cron", "Daily scraping pipeline"],
        ],
        col_widths=[3.5, 6.5, 7]
    )

    add_body(doc, "Technology Stack Details:", bold=True)
    add_bullet(doc, " React 19, Vite 7, Bootstrap 5, Socket.io-client, Recharts, React Router 7", bold_prefix="Frontend:")
    add_bullet(doc, " Node.js, Express 5, Mongoose 9, Puppeteer 24, Cheerio, Helmet, JWT", bold_prefix="Backend:")
    add_bullet(doc, " RAG with text-embedding-3-small (1536D) → Elasticsearch kNN → GPT-4o-mini (temp 0.1)", bold_prefix="AI Pipeline:")
    add_bullet(doc, " bcrypt (12 rounds), JWT HS256 (25min), rate limiting, NoSQL injection guard, prompt injection detection", bold_prefix="Security:")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 3. PROPOSED GCP ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Proposed GCP Architecture", level=1)
    add_thin_line(doc)
    add_body(doc,
        "The proposed architecture leverages GCP-native and GCP-hosted services to provide a fully managed, "
        "auto-scaling, and cost-efficient infrastructure. All services are deployed in the europe-west1 "
        "(Belgium) region for low latency to UNIDO headquarters in Vienna."
    )

    add_styled_table(doc,
        ["Layer", "GCP Service", "Purpose"],
        [
            ["CDN / Frontend", "Firebase Hosting + Cloud CDN", "Serve React SPA globally with caching"],
            ["Security", "Cloud Armor", "WAF, DDoS protection, IP allowlisting"],
            ["Compute", "Cloud Run (fully managed)", "Node.js backend — auto-scales, supports WebSockets"],
            ["AI / LLM", "Vertex AI (Gemini 2.0 Flash)", "Chat completions, response generation"],
            ["Embeddings", "Vertex AI (text-embedding-005)", "Query vectorization for semantic search"],
            ["Vector Search", "Elastic Cloud on GCP", "kNN vector search — zero code changes"],
            ["Database", "MongoDB Atlas on GCP", "Sessions, logs, settings — zero code changes"],
            ["Storage", "Cloud Storage (GCS)", "Scraped data backups with 7-day lifecycle"],
            ["Scheduling", "Cloud Scheduler", "Trigger daily scraping pipeline via HTTP"],
            ["Secrets", "Secret Manager", "API keys, JWT secrets, DB credentials"],
            ["Monitoring", "Cloud Logging + Monitoring", "Logs, uptime checks, alerts, dashboards"],
        ],
        col_widths=[3, 5.5, 8.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 4. GCP SERVICE MAPPING
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. GCP Service Mapping — Detailed", level=1)
    add_thin_line(doc)

    # 4.1
    add_heading_styled(doc, "4.1 Compute — Cloud Run", level=2)
    add_styled_table(doc,
        ["Aspect", "Details"],
        [
            ["Service", "Cloud Run (fully managed, serverless containers)"],
            ["Use Case", "Host Node.js backend (Express 5 + Socket.io)"],
            ["Why Cloud Run", "Auto-scales to zero, supports WebSockets, container-based, no infra management"],
            ["Configuration", "Min instances: 1, Max instances: 10"],
            ["CPU / Memory", "1 vCPU, 1 GB RAM (scalable on demand)"],
            ["Region", "europe-west1 (Belgium) — closest to UNIDO HQ in Vienna"],
        ],
        col_widths=[4, 13]
    )

    # 4.2
    add_heading_styled(doc, "4.2 AI/ML — Vertex AI", level=2)
    add_body(doc, "Three options are available for the AI layer. Each is compatible with the existing RAG architecture:")
    add_styled_table(doc,
        ["Option", "Model", "Pricing (per 1M tokens)", "Pros"],
        [
            ["A (Recommended)", "Gemini 2.0 Flash", "Input: $0.10 / Output: $0.40", "Cheapest, lowest latency on GCP, excellent for RAG"],
            ["B", "OpenAI API Direct (GPT-4o-mini)", "Input: $0.15 / Output: $0.60", "Minimal code changes, proven model quality"],
            ["C", "GPT-4o-mini via Vertex AI", "Input: $0.15 / Output: $0.60 + fee", "All traffic through GCP, unified billing"],
        ],
        col_widths=[3, 5, 4.5, 5.5]
    )
    add_body(doc,
        "Embeddings: Vertex AI text-embedding-005 (768D) or continue with OpenAI text-embedding-3-small (1536D). "
        "If switching to Google embeddings, a one-time re-indexing of Elasticsearch is required.",
        italic=True, color=TEXT_SECONDARY
    )

    # 4.3
    add_heading_styled(doc, "4.3 Vector Search — Elasticsearch on Elastic Cloud", level=2)
    add_styled_table(doc,
        ["Aspect", "Details"],
        [
            ["Service", "Elastic Cloud on GCP (managed Elasticsearch)"],
            ["Why", "Zero code changes — same @elastic/elasticsearch client library, same kNN queries"],
            ["Alternative", "Vertex AI Vector Search (requires code rewrite, saves ~$60/month)"],
            ["Configuration", "1 node, 4 GB RAM, 120 GB storage"],
            ["Region", "GCP europe-west1"],
        ],
        col_widths=[4, 13]
    )

    # 4.4
    add_heading_styled(doc, "4.4 Database — MongoDB Atlas on GCP", level=2)
    add_styled_table(doc,
        ["Aspect", "Details"],
        [
            ["Service", "MongoDB Atlas (deployed on GCP infrastructure)"],
            ["Why", "Zero code changes — Mongoose ORM works with same connection string"],
            ["Configuration", "M10 cluster (2 GB RAM, 10 GB storage) — sufficient for chat logs"],
            ["Region", "GCP europe-west1, VPC peering for private connectivity"],
            ["Collections", "chatSession, chatLog, adminUser, adminSettings"],
        ],
        col_widths=[4, 13]
    )

    # 4.5
    add_heading_styled(doc, "4.5 Object Storage — Cloud Storage (GCS)", level=2)
    add_styled_table(doc,
        ["Aspect", "Details"],
        [
            ["Service", "Google Cloud Storage"],
            ["Replaces", "Azure Blob Storage"],
            ["Use Case", "Store scraped data backups (JSON files)"],
            ["Storage Class", "Standard (or Nearline for cost savings on infrequent access)"],
            ["Lifecycle", "Auto-delete after 7 days (matching current retention policy)"],
        ],
        col_widths=[4, 13]
    )

    # 4.6
    add_heading_styled(doc, "4.6 Frontend Hosting — Firebase Hosting", level=2)
    add_styled_table(doc,
        ["Aspect", "Details"],
        [
            ["Service", "Firebase Hosting (free tier available)"],
            ["Use Case", "Serve React SPA (static build from Vite)"],
            ["Features", "Global CDN, automatic SSL, custom domains, easy CI/CD"],
            ["Alternative", "Cloud Storage + Cloud CDN (more control, slightly more setup)"],
        ],
        col_widths=[4, 13]
    )

    # 4.7
    add_heading_styled(doc, "4.7 Supporting Services", level=2)
    add_styled_table(doc,
        ["Service", "Purpose", "Replaces"],
        [
            ["Cloud Scheduler", "Trigger daily scraping pipeline at 00:00 UTC via HTTP", "node-cron (in-process)"],
            ["Secret Manager", "Store API keys, JWT secrets, DB credentials securely", ".env files"],
            ["Cloud Armor", "WAF, DDoS protection, IP allowlisting", "Application-level only"],
            ["IAM", "Service account permissions, least-privilege access", "N/A"],
            ["Cloud Logging", "Centralized application logs", "Console logging"],
            ["Cloud Monitoring", "Uptime checks, alerts, performance dashboards", "Manual monitoring"],
            ["Error Reporting", "Automatic error detection and grouping", "Manual log review"],
        ],
        col_widths=[3.5, 8.5, 5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 5. COST BREAKDOWN
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Infrastructure Cost Breakdown", level=1)
    add_thin_line(doc)

    # 5.1 One-time
    add_heading_styled(doc, "5.1 One-Time Costs (Setup & Migration)", level=2)
    add_styled_table(doc,
        ["Item", "Cost", "Notes"],
        [
            ["GCP Project setup & IAM configuration", "$0", "Free — console setup"],
            ["Domain DNS migration", "$0", "Free — DNS record update"],
            ["Docker containerization for Cloud Run", "$0", "Development effort only"],
            ["Elasticsearch data migration", "~$50", "Data transfer + re-indexing compute time"],
            ["MongoDB Atlas migration", "~$50", "mongodump/mongorestore + data transfer"],
            ["SSL certificates", "$0", "Managed automatically by GCP/Firebase"],
            ["Cloud Storage bucket setup", "$0", "Free — lifecycle rules configuration"],
            ["Firebase Hosting setup", "$0", "Free — one-time CLI setup"],
            ["Testing & validation environment", "$50 – $150", "Temporary compute for parallel testing"],
            ["Total One-Time Cost", "$150 – $250", ""],
        ],
        col_widths=[7, 2.5, 7.5]
    )

    # 5.2 Monthly — Option A
    add_heading_styled(doc, "5.2 Monthly Recurring Costs", level=2)
    add_body(doc, "Option A: Gemini 2.0 Flash (Recommended — Lowest Cost)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Service", "Configuration", "Monthly Cost"],
        [
            ["Cloud Run (Backend)", "1 vCPU, 1 GB RAM, min 1 instance, ~720 hrs/mo", "$40 – $65"],
            ["Vertex AI — Gemini 2.0 Flash", "~500K input + ~200K output tokens/month", "$10 – $25"],
            ["Vertex AI — Embeddings", "text-embedding-005, ~100K tokens/month", "$1 – $5"],
            ["Elasticsearch (Elastic Cloud)", "1 node, 4 GB RAM, GCP-hosted", "$95 – $120"],
            ["MongoDB Atlas (GCP)", "M10 Dedicated Cluster", "$57 – $80"],
            ["Cloud Storage", "~5 GB Standard, 7-day lifecycle", "$1 – $2"],
            ["Firebase Hosting", "Free tier (10 GB transfer/month)", "$0"],
            ["Cloud Scheduler", "3 scheduled jobs", "$0.30"],
            ["Secret Manager", "~15 secrets, 10K access ops", "$0.60"],
            ["Cloud Logging & Monitoring", "First 50 GB free", "$0 – $10"],
            ["Cloud Armor", "Basic security policy", "$5 – $10"],
            ["Network Egress", "~10 GB/month", "$1 – $5"],
            ["Total Monthly (Option A)", "", "$210 – $325"],
        ],
        col_widths=[5.5, 7, 4.5]
    )

    # Option B
    add_body(doc, "Option B: OpenAI API Direct (GPT-4o-mini)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Service", "Configuration", "Monthly Cost"],
        [
            ["Cloud Run (Backend)", "1 vCPU, 1 GB RAM, min 1 instance", "$40 – $65"],
            ["OpenAI API — GPT-4o-mini", "~500K input + ~200K output tokens/month", "$15 – $40"],
            ["OpenAI API — Embeddings", "text-embedding-3-small, ~100K tokens/month", "$2 – $5"],
            ["Elasticsearch (Elastic Cloud)", "1 node, 4 GB RAM", "$95 – $120"],
            ["MongoDB Atlas (GCP)", "M10 Dedicated Cluster", "$57 – $80"],
            ["Other GCP services", "Storage, CDN, Scheduler, Secrets, Armor, Logging", "$12 – $30"],
            ["Total Monthly (Option B)", "", "$220 – $340"],
        ],
        col_widths=[5.5, 7, 4.5]
    )

    # Option C
    add_body(doc, "Option C: OpenAI on Vertex AI (GPT-4o-mini via Model Garden)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Service", "Monthly Cost"],
        [
            ["All GCP infra (same as above)", "$200 – $300"],
            ["Vertex AI — GPT-4o-mini via Model Garden", "$20 – $50"],
            ["Total Monthly (Option C)", "$240 – $360"],
        ],
        col_widths=[11, 6]
    )

    # 5.3 Comparison
    add_heading_styled(doc, "5.3 Cost Comparison Summary", level=2)
    add_styled_table(doc,
        ["Scenario", "Monthly Cost", "Annual Cost"],
        [
            ["Option A: Gemini Flash (Recommended)", "$210 – $325", "$2,520 – $3,900"],
            ["Option B: OpenAI Direct", "$220 – $340", "$2,640 – $4,080"],
            ["Option C: OpenAI on Vertex", "$240 – $360", "$2,880 – $4,320"],
        ],
        col_widths=[7, 5, 5]
    )

    # 5.4 Optimizations
    add_heading_styled(doc, "5.4 Cost Optimization Opportunities", level=2)
    add_styled_table(doc,
        ["Optimization", "Potential Savings"],
        [
            ["Committed Use Discounts (CUDs) for Cloud Run", "17–40% on compute costs"],
            ["MongoDB Atlas Serverless (M0/M2) instead of M10", "Save ~$50/month"],
            ["Vertex AI Vector Search instead of Elastic Cloud", "Save ~$60/month (requires code changes)"],
            ["Scale Cloud Run to zero during off-hours", "Save ~$15/month"],
            ["Cloud Run Jobs for scraping (pay only when running)", "Save ~$10/month"],
        ],
        col_widths=[10, 7]
    )

    add_highlight_box(doc, "With all optimizations applied: ~$130 – $220/month", bg_color=GREEN_BG)

    # 5.5 Free tier
    add_heading_styled(doc, "5.5 GCP Free Tier Benefits", level=2)
    add_styled_table(doc,
        ["Service", "Free Allowance"],
        [
            ["Cloud Run", "2M requests/month, 360K vCPU-seconds, 180K GiB-seconds"],
            ["Cloud Storage", "5 GB Standard storage"],
            ["Firebase Hosting", "10 GB transfer/month, 1 GB storage"],
            ["Cloud Scheduler", "3 free jobs"],
            ["Secret Manager", "6 active secret versions, 10K access operations"],
            ["Cloud Logging", "50 GB/month"],
            ["Cloud Build", "120 build-minutes/day"],
        ],
        col_widths=[5, 12]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 6. CODE CHANGES
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Required Code Changes", level=1)
    add_thin_line(doc)

    add_heading_styled(doc, "6.1 Changes Summary", level=2)
    add_styled_table(doc,
        ["File / Module", "Change Type", "Effort"],
        [
            ["src/config/openai.js", "Replace Azure OpenAI SDK with Vertex AI / OpenAI SDK", "Medium"],
            ["src/services/embedding.service.js", "Update embedding API calls", "Low"],
            ["src/services/generation.service.js", "Update chat completion API calls", "Low"],
            ["src/config/azureBlob.js", "Replace with GCS client library", "Medium"],
            ["src/storage/azure.storage.js", "Rewrite for Cloud Storage", "Medium"],
            ["src/cron/cronjob.js", "Adapt for Cloud Scheduler HTTP trigger", "Low"],
            ["Dockerfile (new)", "Create container image for Cloud Run", "Low"],
            [".env / Secret Manager", "Update all environment variables", "Low"],
            ["Frontend/src/config/appConfig.js", "Update API URL to Cloud Run endpoint", "Trivial"],
            ["firebase.json (new)", "Firebase Hosting configuration", "Low"],
            ["cloudbuild.yaml (new)", "CI/CD pipeline for Cloud Build", "Low"],
        ],
        col_widths=[6, 8, 3]
    )

    add_heading_styled(doc, "6.2 Detailed Code Changes", level=2)

    add_body(doc, "A. AI Service Migration (Option A — Gemini)", bold=True, color=PRIMARY_DARK)
    add_body(doc, "File: back-end/src/config/openai.js — Replace Azure OpenAI SDK with Vertex AI:")
    add_code_block(doc,
        "// Before (Azure OpenAI)\n"
        "import { AzureOpenAI } from 'openai';\n"
        "const client = new AzureOpenAI({ endpoint, apiKey, apiVersion });\n\n"
        "// After (Vertex AI - Gemini)\n"
        "import { VertexAI } from '@google-cloud/vertexai';\n"
        "const vertexAI = new VertexAI({ project: GCP_PROJECT_ID, location: GCP_REGION });\n"
        "const model = vertexAI.getGenerativeModel({ model: 'gemini-2.0-flash' });"
    )

    add_body(doc, "File: back-end/src/services/generation.service.js — Update completion calls:")
    add_code_block(doc,
        "// Before: client.chat.completions.create({ model, messages, temperature })\n"
        "// After:  model.generateContent({ contents, generationConfig: { temperature: 0.1 } })"
    )

    add_body(doc, "B. Storage Migration", bold=True, color=PRIMARY_DARK)
    add_body(doc, "File: back-end/src/storage/azure.storage.js → gcs.storage.js")
    add_code_block(doc,
        "// Before: import { BlobServiceClient } from '@azure/storage-blob';\n"
        "// After:  import { Storage } from '@google-cloud/storage';\n"
        "//         const bucket = new Storage().bucket(process.env.GCS_BUCKET_NAME);\n"
        "//         await bucket.file(filename).save(data);"
    )

    add_body(doc, "C. Containerization (New Files)", bold=True, color=PRIMARY_DARK)
    add_body(doc, "Create back-end/Dockerfile for Cloud Run deployment:")
    add_code_block(doc,
        "FROM node:20-slim\n"
        "RUN apt-get update && apt-get install -y chromium\n"
        "ENV PUPPETEER_EXECUTABLE_PATH=/usr/bin/chromium\n"
        "WORKDIR /app\n"
        "COPY package*.json ./\n"
        "RUN npm ci --production\n"
        "COPY . .\n"
        "EXPOSE 5000\n"
        'CMD ["node", "src/server.js"]'
    )

    add_body(doc, "D. Cloud Scheduler Integration", bold=True, color=PRIMARY_DARK)
    add_body(doc, "Replace in-process node-cron with an HTTP endpoint triggered by Cloud Scheduler:")
    add_code_block(doc,
        "// Add to server.js routes:\n"
        "app.post('/api/internal/trigger-scrape', authMiddleware, async (req, res) => {\n"
        "  await runPipeline();\n"
        "  res.status(200).json({ success: true });\n"
        "});"
    )

    add_body(doc, "E. New NPM Dependencies", bold=True, color=PRIMARY_DARK)
    add_code_block(doc,
        "// Remove:  @azure/storage-blob\n"
        "// Add (Option A):  @google-cloud/vertexai, @google-cloud/storage\n"
        "// Add (Option B):  @google-cloud/storage  (openai package already installed)"
    )

    add_heading_styled(doc, "6.3 Components Requiring No Changes", level=2)
    add_styled_table(doc,
        ["Component", "Reason"],
        [
            ["Elasticsearch client (@elastic/elasticsearch)", "Same client works with Elastic Cloud on GCP"],
            ["MongoDB / Mongoose ORM", "MongoDB Atlas works on any cloud; same connection string"],
            ["Socket.io (WebSocket)", "Cloud Run supports WebSockets natively"],
            ["Express routes, controllers, middleware", "No cloud-specific code"],
            ["Security middleware (Helmet, rate-limiting)", "Framework-level, fully cloud-agnostic"],
            ["Frontend React app", "Static build deploys anywhere"],
            ["Puppeteer scraper", "Works in Docker container with Chromium"],
            ["Guard service (prompt injection detection)", "Pure application logic, no cloud dependency"],
        ],
        col_widths=[7, 10]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 7. MIGRATION TIMELINE
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Migration Timeline", level=1)
    add_thin_line(doc)
    add_body(doc,
        "The migration follows a four-phase approach, designed for zero-downtime transition. "
        "Total estimated effort: 5 – 8 working days."
    )

    add_body(doc, "Phase 1: Infrastructure Setup (Day 1 – 2)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Task", "Duration", "Details"],
        [
            ["Create GCP project & enable APIs", "2 hours", "Cloud Run, Vertex AI, Storage, Secrets, Scheduler"],
            ["Set up IAM & service accounts", "2 hours", "Least-privilege service account permissions"],
            ["Deploy MongoDB Atlas on GCP", "2 hours", "M10 cluster in europe-west1"],
            ["Set up Elastic Cloud on GCP", "2 hours", "1-node cluster with kNN enabled"],
            ["Create Cloud Storage bucket", "30 min", "Lifecycle rules for 7-day retention"],
            ["Configure Secret Manager", "1 hour", "Store all environment variables securely"],
            ["Set up Firebase Hosting", "30 min", "Connect to custom domain"],
        ],
        col_widths=[6, 2.5, 8.5]
    )

    add_body(doc, "Phase 2: Code Changes (Day 2 – 5)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Task", "Duration", "Details"],
        [
            ["AI service migration (Vertex AI / OpenAI)", "4 – 6 hours", "Config + embedding + generation services"],
            ["Storage migration (Azure Blob → GCS)", "2 – 3 hours", "New storage service implementation"],
            ["Dockerize backend", "2 – 3 hours", "Dockerfile + local testing + optimization"],
            ["Cloud Scheduler integration", "1 – 2 hours", "HTTP trigger endpoint for scraping"],
            ["Environment variable migration", "1 hour", "Secret Manager integration"],
            ["Frontend config update", "30 min", "Update VITE_API_BASE_URL"],
        ],
        col_widths=[6, 2.5, 8.5]
    )

    add_body(doc, "Phase 3: Testing & Validation (Day 5 – 7)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Task", "Duration", "Details"],
        [
            ["Deploy to Cloud Run (staging)", "2 hours", "First deployment + debugging"],
            ["Test RAG pipeline end-to-end", "3 – 4 hours", "Verify embedding → search → generation"],
            ["Test scraping pipeline", "2 hours", "Run full scrape cycle on GCP"],
            ["Test WebSocket (admin dashboard)", "1 hour", "Verify real-time updates via Socket.io"],
            ["Load testing", "2 – 3 hours", "Verify auto-scaling behavior"],
            ["Deploy frontend to Firebase", "1 hour", "Verify CDN, routing, and SSL"],
        ],
        col_widths=[6, 2.5, 8.5]
    )

    add_body(doc, "Phase 4: Go Live (Day 7 – 8)", bold=True, color=PRIMARY_DARK)
    add_styled_table(doc,
        ["Task", "Duration", "Details"],
        [
            ["DNS cutover", "1 hour", "Point domain to Firebase / Cloud Run"],
            ["Monitor for 24 hours", "Ongoing", "Watch logs, errors, latency metrics"],
            ["Decommission old infrastructure", "1 hour", "After successful validation period"],
        ],
        col_widths=[6, 2.5, 8.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 8. ARCHITECTURE DIAGRAMS
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Architecture Diagrams", level=1)
    add_thin_line(doc)

    add_heading_styled(doc, "8.1 Data Flow", level=2)
    add_styled_table(doc,
        ["Step", "Component", "Action"],
        [
            ["1", "User (Browser)", "Sends question via chat widget"],
            ["2", "Firebase Hosting", "Serves React SPA, routes API calls to Cloud Run"],
            ["3", "Cloud Run (Backend)", "Receives question, calls Vertex AI for embedding"],
            ["4", "Vertex AI (Embeddings)", "Converts question to 768D/1536D vector"],
            ["5", "Elasticsearch", "kNN search: finds top 5 relevant document chunks"],
            ["6", "Cloud Run (Backend)", "Filters by relevance score (≥0.75), builds prompt"],
            ["7", "Vertex AI (Gemini Flash)", "Generates answer from context + question (temp 0.1)"],
            ["8", "MongoDB Atlas", "Logs session, question, answer, metadata"],
            ["9", "User (Browser)", "Receives and displays answer"],
        ],
        col_widths=[1.5, 4.5, 11]
    )

    add_heading_styled(doc, "8.2 Security Architecture", level=2)
    add_styled_table(doc,
        ["Layer", "Service", "Protection"],
        [
            ["Edge", "Cloud Armor", "WAF rules, DDoS mitigation, geo-blocking"],
            ["Transport", "Cloud Load Balancer", "TLS 1.3 termination, automatic SSL certificates"],
            ["Application", "Express Middleware", "Helmet, rate limiting, CORS, input sanitization, NoSQL injection guard"],
            ["AI Safety", "Guard Service", "Prompt injection detection, content filtering"],
            ["Auth", "JWT + bcrypt", "HS256 tokens (25min), bcrypt-12 password hashing"],
            ["Network", "VPC Peering", "Private connectivity: Cloud Run ↔ MongoDB ↔ Elasticsearch"],
            ["Secrets", "Secret Manager", "IAM-controlled access, automatic rotation support"],
        ],
        col_widths=[2.5, 4.5, 10]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 9. RECOMMENDATIONS
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Recommendations", level=1)
    add_thin_line(doc)

    add_heading_styled(doc, "9.1 Recommended Option: Option A (Gemini 2.0 Flash)", level=2)
    add_styled_table(doc,
        ["Factor", "Reasoning"],
        [
            ["Cost", "30–60% cheaper than GPT-4o-mini for comparable quality on RAG tasks"],
            ["Latency", "Lower latency when both AI and infrastructure are on GCP"],
            ["Data Residency", "All data stays within GCP Europe region"],
            ["Vendor Lock-in", "Moderate — can switch to OpenAI later with config change"],
            ["Quality", "Gemini 2.0 Flash is excellent for RAG tasks at this scale"],
        ],
        col_widths=[3, 14]
    )

    add_heading_styled(doc, "9.2 If OpenAI Quality is Preferred: Option B", level=2)
    add_bullet(doc, "Minimal code changes — just remove Azure wrapper and use standard OpenAI SDK")
    add_bullet(doc, "Same model quality as current setup (GPT-4o-mini)")
    add_bullet(doc, "Slightly higher AI costs but lower migration effort")

    add_heading_styled(doc, "9.3 Additional Recommendations", level=2)
    add_bullet(doc, " Free tier or $9/month is sufficient for initial load; upgrade later as needed", bold_prefix="1. Start with MongoDB Atlas Shared (M0/M2) —")
    add_bullet(doc, " Ensures WebSocket connections stay alive without cold starts", bold_prefix="2. Use Cloud Run min-instances=1 —")
    add_bullet(doc, " on the Cloud Run service for API response caching on GET endpoints", bold_prefix="3. Enable Cloud CDN")
    add_bullet(doc, " for CI/CD — auto-deploy on git push to main branch", bold_prefix="4. Set up Cloud Build")
    add_bullet(doc, " as future replacement for Elasticsearch — saves ~$60–100/month but requires code rewrite", bold_prefix="5. Consider Vertex AI Vector Search")
    add_bullet(doc, " with Cloud Logging client for better observability", bold_prefix="6. Implement structured logging")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # APPENDIX A
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Appendix A: Environment Variables (GCP)", level=1)
    add_thin_line(doc)
    add_styled_table(doc,
        ["Variable", "Required", "Default", "Purpose"],
        [
            ["GCP_PROJECT_ID", "✓", "—", "GCP project identifier"],
            ["GCP_REGION", "✓", "europe-west1", "GCP deployment region"],
            ["VERTEX_AI_MODEL", "✓ (Opt A)", "gemini-2.0-flash", "Chat model deployment"],
            ["VERTEX_EMBEDDING_MODEL", "✓ (Opt A)", "text-embedding-005", "Embedding model"],
            ["OPENAI_API_KEY", "✓ (Opt B)", "—", "OpenAI API key (if using Option B)"],
            ["MONGO_URI", "✓", "—", "MongoDB Atlas connection string"],
            ["ELASTIC_NODE", "✓", "—", "Elasticsearch endpoint URL"],
            ["ELASTIC_USERNAME", "○", "—", "Elasticsearch auth username"],
            ["ELASTIC_PASSWORD", "○", "—", "Elasticsearch auth password"],
            ["GCS_BUCKET_NAME", "✓", "—", "Cloud Storage bucket for backups"],
            ["JWT_SECRET", "✓", "—", "JWT signing key (32+ chars)"],
            ["ADMIN_EMAIL", "✓", "—", "Default admin email"],
            ["ADMIN_PASSWORD", "✓", "—", "Default admin password"],
            ["NODE_ENV", "○", "production", "Environment mode"],
            ["PORT", "○", "5000", "Backend server port"],
            ["ENABLE_CRON", "○", "false", "Enable in-process cron"],
            ["RAG_MIN_SCORE", "○", "0.75", "Minimum relevance score for search"],
            ["RAG_TOP_K", "○", "5", "Number of documents retrieved"],
        ],
        col_widths=[5, 2, 3.5, 6.5]
    )
    add_body(doc, "✓ = Required    ○ = Optional", italic=True, color=TEXT_SECONDARY)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # APPENDIX B
    # ════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Appendix B: GCP CLI Setup Commands", level=1)
    add_thin_line(doc)

    add_body(doc, "1. Create project and enable APIs:", bold=True)
    add_code_block(doc,
        "gcloud projects create unido-chatbot-prod\n"
        "gcloud services enable run.googleapis.com aiplatform.googleapis.com \\\n"
        "  storage.googleapis.com secretmanager.googleapis.com \\\n"
        "  cloudscheduler.googleapis.com cloudbuild.googleapis.com \\\n"
        "  firebase.googleapis.com"
    )

    add_body(doc, "2. Deploy backend to Cloud Run:", bold=True)
    add_code_block(doc,
        "gcloud run deploy unido-chatbot --source . \\\n"
        "  --region europe-west1 --min-instances 1 --max-instances 10 \\\n"
        "  --memory 1Gi --cpu 1 --port 5000 --allow-unauthenticated"
    )

    add_body(doc, "3. Create Cloud Scheduler job:", bold=True)
    add_code_block(doc,
        "gcloud scheduler jobs create http daily-scrape \\\n"
        "  --schedule='0 0 * * *' \\\n"
        "  --uri='https://unido-chatbot-xxxxx.run.app/api/internal/trigger-scrape' \\\n"
        "  --http-method=POST \\\n"
        "  --oidc-service-account-email=scheduler-sa@unido-chatbot-prod.iam.gserviceaccount.com"
    )

    add_body(doc, "4. Create Cloud Storage bucket:", bold=True)
    add_code_block(doc,
        "gsutil mb -l europe-west1 gs://unido-chatbot-backups\n"
        "gsutil lifecycle set lifecycle.json gs://unido-chatbot-backups"
    )

    # End of document
    doc.add_paragraph()
    p_line = doc.add_paragraph()
    run = p_line.add_run("━" * 80)
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(6)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_end.add_run("— End of Document —")
    run.font.color.rgb = TEXT_SECONDARY
    run.font.size = Pt(10)
    run.italic = True

    # Save
    output_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'GCP_Infrastructure_Proposal.docx')
    output_path = os.path.abspath(output_path)
    doc.save(output_path)
    print(f"Word document generated successfully: {output_path}")
    return output_path


if __name__ == "__main__":
    build_docx()
