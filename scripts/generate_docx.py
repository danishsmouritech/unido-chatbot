"""
Generate a professional Word document for the UNIDO RAG Chatbot Technical Documentation.
Mirrors the 16-page MouriTech-branded PDF.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colours ──
BLUE_DARK  = RGBColor(0x0A, 0x1F, 0x3D)
BLUE_MID   = RGBColor(0x00, 0x66, 0xB3)
BLUE_LIGHT = RGBColor(0x33, 0x7A, 0xB7)
GREEN      = RGBColor(0x05, 0x96, 0x69)
ORANGE     = RGBColor(0xE8, 0x8C, 0x30)
RED_DARK   = RGBColor(0xC0, 0x39, 0x2B)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_DARK  = RGBColor(0x33, 0x41, 0x55)
GREY_MID   = RGBColor(0x64, 0x74, 0x8B)
GREY_LIGHT = RGBColor(0xF1, 0xF5, 0xF9)
BLACK      = RGBColor(0x1E, 0x29, 0x3B)

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = GREY_DARK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15


def set_cell_shading(cell, color_hex):
    """Set cell background colour."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val["sz"]}" '
                f'w:space="0" w:color="{val["color"]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_line(color="0066B3", thickness=6):
    """Add a horizontal line via a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, color)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pf = p._element
        pPr = pf.get_or_add_pPr()
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="40" w:lineRule="exact"/>')
        pPr.append(spacing)
    # set row height
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="40" w:hRule="exact"/>')
    trPr.append(trHeight)
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    return tbl


def remove_table_borders(tbl):
    """Remove all borders from a table."""
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_section_heading(number, title):
    """Add a styled section heading with blue left-bar accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    # Blue bar character + heading
    run_bar = p.add_run("█  ")
    run_bar.font.color.rgb = BLUE_MID
    run_bar.font.size = Pt(18)
    run = p.add_run(f"{number}. {title}")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = GREY_DARK
    return p


def add_sub_heading(number, title, color=BLUE_MID):
    """Add a styled sub-heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{number}  {title}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_body(text):
    """Add body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = GREY_DARK
    return p


def add_bullet(text, bold_prefix=None, color=BLUE_MID):
    """Add a styled bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0)
    # Bullet dot
    run_dot = p.add_run("●  ")
    run_dot.font.size = Pt(7)
    run_dot.font.color.rgb = color
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.color.rgb = GREY_DARK
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = GREY_DARK
    return p


def add_styled_table(headers, rows, col_widths=None, header_color="33537E"):
    """Add a professional styled table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, header_color)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = GREY_DARK
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_info_card(badge, badge_color, title, description, bullets, bullet_color=BLUE_MID):
    """Add an infrastructure-style info card."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    # Badge column
    cell_badge = tbl.cell(0, 0)
    cell_badge.width = Cm(1.5)
    p = cell_badge.paragraphs[0]
    run = p.add_run(badge)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell_badge, badge_color)
    set_cell_borders(cell_badge, left={"sz": "12", "color": badge_color})
    # Content column
    cell_content = tbl.cell(0, 1)
    cell_content.text = ""
    p_title = cell_content.paragraphs[0]
    run_t = p_title.add_run(title)
    run_t.font.bold = True
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = GREY_DARK
    p_desc = cell_content.add_paragraph()
    run_d = p_desc.add_run(description)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = GREY_MID
    for b in bullets:
        p_b = cell_content.add_paragraph()
        r_dot = p_b.add_run("●  ")
        r_dot.font.size = Pt(7)
        r_dot.font.color.rgb = bullet_color
        r_text = p_b.add_run(b)
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = GREY_DARK
        p_b.paragraph_format.left_indent = Cm(0.5)
        p_b.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()  # spacer


def add_page_footer_text(text):
    """Add a footer-like text at the bottom."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line("0066B3", 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY_MID
    run.font.italic = True


def add_page_break():
    doc.add_page_break()


def add_stat_boxes(stats):
    """Add KPI/stat boxes in a row."""
    tbl = doc.add_table(rows=2, cols=len(stats))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    colors = ["33537E", "E88C30", "059669", "C0392B"]
    for i, (value, label) in enumerate(stats):
        color = colors[i % len(colors)]
        # Value cell
        cell_val = tbl.rows[0].cells[i]
        cell_val.text = ""
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color
        set_cell_shading(cell_val, "F8FAFC")
        set_cell_borders(cell_val, top={"sz": "8", "color": color})
        # Label cell
        cell_lbl = tbl.rows[1].cells[i]
        cell_lbl.text = ""
        p2 = cell_lbl.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(8.5)
        run2.font.color.rgb = GREY_MID
        set_cell_shading(cell_lbl, "F8FAFC")
    doc.add_paragraph()


def add_flow_steps(steps):
    """Add a horizontal flow diagram using a table."""
    # steps is a list of (number, label) tuples
    cols = len(steps)
    tbl = doc.add_table(rows=2, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    colors = ["33537E", "C0392B", "7C3AED", "E88C30", "059669", "33537E", "0066B3"]
    for i, (num, label) in enumerate(steps):
        color = colors[i % len(colors)]
        # Number row
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(num))
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, color)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Label row
        cell2 = tbl.rows[1].cells[i]
        cell2.text = ""
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.bold = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = GREY_DARK
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


def add_deployment_phase(title, badge_text, badge_color, items):
    """Add a deployment phase card."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    # Content
    cell = tbl.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(badge_color)
    set_cell_borders(cell, left={"sz": "12", "color": badge_color})
    for item in items:
        p_b = cell.add_paragraph()
        r_dot = p_b.add_run("●  ")
        r_dot.font.size = Pt(7)
        r_dot.font.color.rgb = RGBColor.from_string(badge_color)
        r_text = p_b.add_run(item)
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = GREY_DARK
        p_b.paragraph_format.left_indent = Cm(0.5)
        p_b.paragraph_format.space_after = Pt(2)
    # Badge
    cell_badge = tbl.cell(0, 1)
    cell_badge.width = Cm(3)
    cell_badge.text = ""
    p_badge = cell_badge.paragraphs[0]
    p_badge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_b = p_badge.add_run(badge_text)
    run_b.font.bold = True
    run_b.font.size = Pt(9)
    run_b.font.color.rgb = WHITE
    set_cell_shading(cell_badge, badge_color)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# MouriTech header
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("MOURITECH  |  GLOBAL ENTERPRISE SOLUTIONS")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = BLUE_MID

add_horizontal_line("0066B3", 8)
add_horizontal_line("E88C30", 4)

doc.add_paragraph()  # spacer

p = doc.add_paragraph()
run = p.add_run("COMPREHENSIVE TECHNICAL DOCUMENTATION")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = BLUE_MID
run.font.all_caps = True

doc.add_paragraph()

# Main title
p = doc.add_paragraph()
run = p.add_run("UNIDO - AI Integration\n& RAG Chatbot")
run.font.bold = True
run.font.size = Pt(36)
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
run = p.add_run(
    "An enterprise-grade, AI-powered recruitment assistant\n"
    "delivering intelligent, context-aware career guidance\n"
    "for UNIDO's global talent ecosystem."
)
run.font.size = Pt(12)
run.font.color.rgb = GREY_DARK

doc.add_paragraph()

# Feature highlights
highlights = [
    ("[AI]  RAG Architecture", "Azure OpenAI + Elasticsearch"),
    ("[SEC]  Enterprise Security", "JWT | Rate Limiting | RBAC"),
    ("[SYS]  Automated Pipeline", "Scrape → Embed → Index → Serve"),
    ("[DATA]  Real-time Analytics", "Live dashboards & audit logs"),
]
tbl = doc.add_table(rows=2, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
for i, (title, desc) in enumerate(highlights):
    cell = tbl.cell(0, i)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = BLUE_MID
    set_cell_shading(cell, "0D2B52")
    cell2 = tbl.cell(1, i)
    cell2.text = ""
    p2 = cell2.paragraphs[0]
    run2 = p2.add_run(desc)
    run2.font.size = Pt(8)
    run2.font.color.rgb = WHITE
    set_cell_shading(cell2, "0D2B52")

doc.add_paragraph()
doc.add_paragraph()

# Metadata table
meta = [
    ("Version", "1.0"),
    ("Date", "April 2026"),
    ("Classification", "Internal - Organizational"),
    ("Prepared by", "MouriTech Global Enterprise Solutions"),
    ("Client", "UNIDO - United Nations Industrial Development Organization"),
]
tbl = doc.add_table(rows=len(meta), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
for i, (key, val) in enumerate(meta):
    tbl.cell(i, 0).text = ""
    p_k = tbl.cell(i, 0).paragraphs[0]
    run_k = p_k.add_run(key)
    run_k.font.size = Pt(10)
    run_k.font.color.rgb = GREY_MID
    tbl.cell(i, 1).text = ""
    p_sep = tbl.cell(i, 1).paragraphs[0]
    run_sep = p_sep.add_run(":")
    run_sep.font.size = Pt(10)
    run_sep.font.color.rgb = GREY_MID
    tbl.cell(i, 2).text = ""
    p_v = tbl.cell(i, 2).paragraphs[0]
    run_v = p_v.add_run(val)
    run_v.font.bold = True
    run_v.font.size = Pt(10)
    run_v.font.color.rgb = GREY_DARK
    if i < len(meta) - 1:
        set_cell_borders(tbl.cell(i, 0), bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_borders(tbl.cell(i, 1), bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_borders(tbl.cell(i, 2), bottom={"sz": "2", "color": "E2E8F0"})

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL  |  MouriTech  |  UNIDO AI Chatbot  |  Version 1.0  |  April 2026")
run.font.size = Pt(8)
run.font.color.rgb = GREY_MID
run.font.bold = True


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Table of Contents")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = GREY_DARK

add_horizontal_line("0066B3", 6)
doc.add_paragraph()

toc_entries = [
    ("1.", "Executive Summary", "03", True),
    ("2.", "System Architecture", "04", True),
    ("  2.1", "Architecture Layers", "04", False),
    ("  2.2", "Simplified Data Flow", "05", False),
    ("3.", "Technology Stack", "06", True),
    ("4.", "Core Components & Capabilities", "07", True),
    ("  4.1", "User-Facing Chatbot Interface", "07", False),
    ("  4.2", "Admin Dashboard", "08", False),
    ("  4.3", "Backend REST API", "08", False),
    ("  4.4", "Automated Data Pipeline", "09", False),
    ("  4.5", "AI & Search Integration", "09", False),
    ("5.", "Security Architecture", "10", True),
    ("  5.1", "Security Model", "10", False),
    ("  5.2", "API Security & Rate Limiting", "11", False),
    ("  5.3", "Input Validation & Sanitisation", "11", False),
    ("  5.4", "Data Protection & Encryption", "12", False),
    ("  5.5", "Guardrails & Content Filtering", "12", False),
    ("6.", "Infrastructure Requirements", "13", True),
    ("7.", "Deployment Strategy", "14", True),
    ("8.", "Analytics & Monitoring", "15", True),
    ("9.", "What We Need From UNIDO to Go Live", "16", True),
    ("10.", "Conclusion", "17", True),
]

tbl = doc.add_table(rows=len(toc_entries), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
for i, (num, title, page, is_main) in enumerate(toc_entries):
    # Number
    cell_num = tbl.cell(i, 0)
    cell_num.width = Cm(1.5)
    cell_num.text = ""
    p = cell_num.paragraphs[0]
    run = p.add_run(num)
    run.font.bold = is_main
    run.font.size = Pt(10.5) if is_main else Pt(10)
    run.font.color.rgb = GREY_DARK if is_main else GREY_MID
    # Title
    cell_title = tbl.cell(i, 1)
    cell_title.text = ""
    p = cell_title.paragraphs[0]
    run = p.add_run(title)
    run.font.bold = is_main
    run.font.size = Pt(10.5) if is_main else Pt(10)
    run.font.color.rgb = GREY_DARK if is_main else GREY_MID
    # Page number
    cell_pg = tbl.cell(i, 2)
    cell_pg.width = Cm(1.5)
    cell_pg.text = ""
    p = cell_pg.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(page)
    run.font.bold = is_main
    run.font.size = Pt(10.5) if is_main else Pt(10)
    run.font.color.rgb = BLUE_MID if is_main else GREY_MID
    # Bottom border for main entries
    if is_main:
        set_cell_borders(cell_num, bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_borders(cell_title, bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_borders(cell_pg, bottom={"sz": "2", "color": "E2E8F0"})


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("1", "Executive Summary")

add_body(
    "The UNIDO RAG Chatbot is a production-ready, AI-powered recruitment assistant built on Microsoft Azure. "
    "It is purpose-built to enhance user engagement on UNIDO's Careers portal by delivering intelligent, "
    "context-aware answers about job opportunities, eligibility criteria, benefits, and application processes "
    "— all grounded in real, verified UNIDO data."
)

doc.add_paragraph()
add_stat_boxes([
    ("95%+", "RAG-grounded\nResponse Accuracy"),
    ("24h", "auto-refresh\nData Freshness"),
    ("100", "req / minute\nRate Limit - Chat"),
    ("5", "req / minute\nRate Limit - Auth"),
])

add_sub_heading("", "What It Delivers", GREY_DARK)
add_bullet("Accurate, grounded answers on jobs, categories, and application processes")
add_bullet("Retrieval-Augmented Generation (RAG) pipeline prevents hallucinations")
add_bullet("Automated ingestion keeps data current without manual intervention")
add_bullet("Real-time WebSocket communication for seamless user experience")
add_bullet("Enterprise-grade security: JWT auth · RBAC · rate limiting · content guardrails")
add_bullet("Full audit trail: every query, response, and administrative action is logged")

doc.add_paragraph()
add_sub_heading("", "Why RAG?", ORANGE)
add_bullet("Traditional chatbots guess — RAG retrieves. Every answer is sourced from verified UNIDO content.", color=ORANGE)
add_bullet("Semantic vector search (Elasticsearch) finds contextually relevant chunks, not just keyword matches.", color=ORANGE)
add_bullet("Azure OpenAI synthesises a natural, readable response from retrieved evidence.", color=ORANGE)
add_bullet("If no relevant data exists, the system responds transparently — no fabricated answers.", color=ORANGE)

doc.add_paragraph()
add_sub_heading("", "Strategic Value", GREY_DARK)
add_body(
    "Deploying the UNIDO RAG Chatbot reduces repetitive HR and recruitment queries, freeing staff for high-value tasks "
    "while providing a 24/7 self-service experience to candidates worldwide. The system is built to scale and adapt as "
    "UNIDO's career offerings evolve."
)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("2", "System Architecture")

add_body(
    "The system follows a microservices-inspired, layered architecture hosted on the Azure cloud platform. "
    "Each layer has a single responsibility, enabling independent scaling, testing, and maintenance."
)

doc.add_paragraph()

# Architecture layers diagram
arch_layers = [
    ("Frontend  |  React + Vite + Socket.io", "0066B3"),
    ("Backend  |  Node.js + Express + Socket.IO Server", "059669"),
    ("AI Layer  |  Azure OpenAI + Elasticsearch Vector Index", "33537E"),
    ("Data  |  MongoDB · Azure Blob Storage", "E88C30"),
    ("Pipeline  |  Node-Cron · Puppeteer Scrapers · Chunk & Embed", "C0392B"),
]
for label, color in arch_layers:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    set_cell_shading(cell, color)
    remove_table_borders(tbl)

doc.add_paragraph()

add_sub_heading("2.1", "Architecture Layers")

layers = [
    ("Presentation Layer", "React + Vite SPA with Socket.io client. Renders the chat widget and admin dashboard. Real-time updates via WebSocket (Socket.IO).", "0066B3"),
    ("Application Layer", "Node.js + Express server. Exposes REST APIs, handles JWT auth, manages Socket.IO events, enforces rate limits, and orchestrates business logic.", "059669"),
    ("AI & Search Layer", "RAG engine: query embedding via Azure OpenAI text-embedding-3-small (1536-dim), semantic k-NN search via Elasticsearch, answer synthesis via Azure OpenAI GPT-4o-mini.", "33537E"),
    ("Data & Storage Layer", "MongoDB stores sessions, conversation history, admin settings, and audit logs. Elasticsearch holds vector embeddings. Azure Blob Storage archives raw scraped JSON.", "E88C30"),
    ("Ingestion / Pipeline Layer", "Node-Cron scheduler triggers Puppeteer-based scrapers on a configurable cadence. Data is cleaned, semantically chunked, embedded, and indexed automatically.", "C0392B"),
]

for name, desc, color in layers:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_borders(cell, left={"sz": "12", "color": color})
    p_name = cell.paragraphs[0]
    run = p_name.add_run(name)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)
    p_desc = cell.add_paragraph()
    run_d = p_desc.add_run(desc)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = GREY_DARK
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — SIMPLIFIED DATA FLOW
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_sub_heading("2.2", "Simplified Data Flow")

add_body(
    "Every user interaction passes through a well-defined, auditable pipeline. "
    "Retrieval-Augmented Generation ensures responses are always grounded in current data."
)

doc.add_paragraph()
add_flow_steps([
    ("1", "Query"), ("2", "Guard"), ("3", "Embed"), ("4", "Search"),
    ("5", "Generate"), ("6", "Log"), ("7", "Deliver")
])

add_sub_heading("", "Flow Step Details", GREY_DARK)

flow_details = [
    ("01", "User Submits Query", "User types a natural-language question into the chat widget. The message is validated (non-empty, within size limit) and a session ID is confirmed."),
    ("02", "Input Validation & Guard Check", "The Guard Service screens the query for offensive language, confidential data requests, or out-of-scope topics. Blocked queries receive a safe fallback and are logged for review."),
    ("03", "Embedding Generation", "The validated query is forwarded to Azure OpenAI (text-embedding-3-small) to produce a 1,536-dimensional vector representation capturing semantic meaning."),
    ("04", "Vector Search — Top-K Retrieval", "Elasticsearch performs an approximate nearest-neighbour search across the career content index, returning the K most contextually relevant document chunks."),
    ("05", "Context-Aware Response Generation", "Retrieved chunks are assembled into a structured prompt and submitted to Azure OpenAI (GPT-4o-mini). The model synthesises a concise, accurate answer."),
    ("06", "Logging & Analytics", "The full interaction — query, retrieved chunks, response, latency, session ID — is persisted in MongoDB for audit, analytics, and continuous quality improvement."),
    ("07", "Response Delivered", "The answer is streamed back to the frontend via Socket.IO and rendered in the chat UI. Relevant source references may be appended for transparency."),
]

colors = ["33537E", "C0392B", "7C3AED", "E88C30", "059669", "33537E", "0066B3"]
for i, (num, title, desc) in enumerate(flow_details):
    color = colors[i % len(colors)]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    # Number badge
    cell_num = tbl.cell(0, 0)
    cell_num.width = Cm(1.5)
    cell_num.text = ""
    p = cell_num.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(num)
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    set_cell_shading(cell_num, color)
    # Content
    cell_content = tbl.cell(0, 1)
    cell_content.text = ""
    p_t = cell_content.paragraphs[0]
    run_t = p_t.add_run(title)
    run_t.font.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor.from_string(color)
    p_d = cell_content.add_paragraph()
    run_d = p_d.add_run(desc)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = GREY_DARK


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("3", "Technology Stack")

add_body(
    "The solution is assembled from best-in-class, enterprise-proven technologies. "
    "Each component was selected for reliability, scalability, and Azure ecosystem compatibility."
)

doc.add_paragraph()

tech_groups = [
    ("Frontend", "33537E", ["React 18 + Vite (build)", "Socket.IO Client (real-time)", "Axios (HTTP)", "Bootstrap / CSS Modules"]),
    ("Backend", "33537E", ["Node.js 20 LTS + Express 5+", "Socket.IO Server", "Mongoose ODM", "express-rate-limit", "jsonwebtoken + bcryptjs"]),
    ("AI & Search", "059669", ["Azure OpenAI (GPT-4o-mini)", "Azure OpenAI Embeddings (text-embedding-3-small)", "Elasticsearch 8.x (kNN 1536-dim)", "LangChain.js (RAG orchestration)"]),
    ("Data Storage", "059669", ["MongoDB Atlas (sessions, logs, admin)", "Elasticsearch Indices (vector store)", "Azure Blob Storage (JSON backups)"]),
    ("Pipeline & Tooling", "E88C30", ["Puppeteer (headless browser scraping)", "Node-Cron (scheduler)", "Custom chunking & embedding scripts"]),
    ("Infrastructure", "C0392B", ["Azure App Service (Node.js runtime)", "Azure Key Vault (secrets management)", "Azure Monitor + Application Insights", "Docker (containerisation)"]),
]

# Layout as 2-column pairs
for i in range(0, len(tech_groups), 2):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    for j in range(2):
        if i + j < len(tech_groups):
            name, color, items = tech_groups[i + j]
            cell = tbl.cell(0, j)
            cell.text = ""
            # Header
            p_h = cell.paragraphs[0]
            run_h = p_h.add_run(f"  {name}")
            run_h.font.bold = True
            run_h.font.size = Pt(11)
            run_h.font.color.rgb = WHITE
            set_cell_shading(cell, color)
            # Items (add as separate paragraph, switch to white bg)
    doc.add_paragraph()
    # Items below each group
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl2)
    for j in range(2):
        if i + j < len(tech_groups):
            name, color, items = tech_groups[i + j]
            cell = tbl2.cell(0, j)
            cell.text = ""
            set_cell_shading(cell, "F8FAFC")
            for item in items:
                p_b = cell.add_paragraph() if cell.paragraphs[0].text == "" and len(cell.paragraphs) > 0 else cell.paragraphs[0]
                if cell.paragraphs[0].text != "":
                    p_b = cell.add_paragraph()
                r_dot = p_b.add_run("●  ")
                r_dot.font.size = Pt(7)
                r_dot.font.color.rgb = RGBColor.from_string(color)
                r_text = p_b.add_run(item)
                r_text.font.size = Pt(10)
                r_text.font.color.rgb = GREY_DARK
                p_b.paragraph_format.space_after = Pt(2)
                p_b.paragraph_format.left_indent = Cm(0.3)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — CORE COMPONENTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("4", "Core Components & Capabilities")

add_sub_heading("4.1", "User-Facing Chatbot Interface")
add_body(
    "The chat widget is embedded directly into the UNIDO Careers website. It provides a conversational interface "
    "modelled on familiar messaging apps, ensuring minimal friction for first-time users."
)
chatbot_features = [
    "Natural language question input with real-time typing indicator",
    "Streaming responses via WebSocket — answers appear word-by-word",
    "Session persistence: conversation context carried across turns",
    "Mobile-responsive layout for on-the-go candidates",
    "Accessible design (WCAG 2.1 AA compliant colour contrast)",
    "Graceful error messages when the chatbot is disabled or offline",
]
for f in chatbot_features:
    add_bullet(f)

doc.add_paragraph()

add_sub_heading("4.2", "Admin Dashboard")
admin_features = [
    "Secure login with JWT — session expires automatically",
    "Live analytics: query volume, top topics, response times",
    "Enable / disable chatbot globally with a single toggle",
    "Trigger manual data scraping & reindexing on demand",
    "Export conversation logs as CSV for compliance and review",
    "User management and role-based access control",
]
for f in admin_features:
    add_bullet(f)

doc.add_paragraph()

add_sub_heading("4.3", "Backend REST API")
add_body(
    "The Express 5 backend exposes RESTful endpoints organised into three route groups: "
    "chat (public), admin (JWT-protected), and auth (rate-limited). All endpoints are documented "
    "via OpenAPI/Swagger."
)

add_sub_heading("4.4", "Automated Data Pipeline")
add_body(
    "A scheduled pipeline (node-cron, default: daily at midnight UTC) scrapes the UNIDO Careers website, "
    "cleans HTML content via Cheerio, splits it into semantic chunks, generates embeddings via Azure OpenAI, "
    "and bulk-indexes everything into Elasticsearch. Backups are saved to local JSON and Azure Blob Storage."
)

add_sub_heading("4.5", "AI & Search Integration")
add_body(
    "The RAG engine combines Azure OpenAI embeddings (text-embedding-3-small, 1536 dimensions) with "
    "Elasticsearch kNN cosine search (top-K=5, min score 0.75). Retrieved chunks are assembled into a structured "
    "prompt and submitted to GPT-4o-mini (temperature 0.1) for near-deterministic, grounded responses."
)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 8-9 — SECURITY ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("5", "Security Architecture")

add_body(
    "The UNIDO RAG Chatbot implements a defence-in-depth security strategy aligned with OWASP Top 10 "
    "mitigations and enterprise cloud best practices."
)

# Defence layers diagram
doc.add_paragraph()
security_layers = [
    ("WAF / DDoS", "C0392B"), ("Rate Limiting", "E88C30"), ("JWT Auth + RBAC", "33537E"),
    ("Input Guard", "059669"), ("Data Encryption", "7C3AED"), ("Audit Logging", "0066B3"),
]
tbl = doc.add_table(rows=1, cols=len(security_layers))
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
for i, (label, color) in enumerate(security_layers):
    cell = tbl.cell(0, i)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = WHITE
    set_cell_shading(cell, color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("← Outer  ─────────────────  Inner →    Each layer independently blocks threats")
run.font.size = Pt(8.5)
run.font.italic = True
run.font.color.rgb = GREY_MID

doc.add_paragraph()

add_sub_heading("5.1", "Security Model")
sec_model = [
    "JWT Bearer token authentication on all protected API endpoints",
    "BCrypt (cost factor 12) password hashing for admin credentials",
    "Role-Based Access Control — admin role enforced at middleware level",
    "Socket.IO authentication middleware: every WS connection validated",
    "Admin kill-switch: chatbot can be disabled globally within seconds",
]
for s in sec_model:
    add_bullet(s)

add_sub_heading("5.2", "API Security & Rate Limiting")
add_body("Tiered rate limiting using express-rate-limit protects against abuse and DoS:")

add_styled_table(
    ["Endpoint Group", "Path Pattern", "Limit", "Purpose"],
    [
        ["Chat Endpoints", "/api/chat/*", "100 req / min", "General users"],
        ["Admin Endpoints", "/api/admin/*", "30 req / min", "Dashboard operators"],
        ["Authentication Endpoints", "/api/auth/*", "5 req / min", "Brute-force protection"],
    ],
    col_widths=[4.5, 3.5, 3, 4.5]
)

add_sub_heading("5.3", "Input Validation & Sanitisation")
input_val = [
    "Request payloads capped at 2 MB (JSON + URL-encoded data)",
    "sessionId and question fields validated as required non-empty strings",
    "Guard Service: pattern-based profanity & sensitive keyword blocking",
    "Automatic safe fallback responses for blocked queries",
    "All blocked queries logged with IP address and timestamp for review",
]
for v in input_val:
    add_bullet(v)

add_sub_heading("5.4", "Data Protection & Encryption")
data_prot = [
    "All secrets (API keys, DB URIs, JWT secret) stored in Azure Key Vault",
    "MongoDB connection strings never hard-coded — injected via environment",
    "Admin passwords encrypted at rest in MongoDB (BCrypt hash stored)",
    "TLS/HTTPS enforced end-to-end in production environments",
    "IP addresses and User-Agent strings captured for full audit trail",
]
for d in data_prot:
    add_bullet(d)

add_sub_heading("5.5", "Guardrails & Content Filtering")
add_body("Content Guard — How It Works:")
guard_steps = [
    "Incoming query → pattern-match against offensive language list → BLOCK if matched",
    "Remaining query → check against confidential/private keyword list → BLOCK if matched",
    "Passed query → proceed to embedding & RAG pipeline",
    "All blocked events written to MongoDB audit collection with full metadata",
]
for g in guard_steps:
    add_bullet(g, color=GREEN)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 10-11 — INFRASTRUCTURE REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("6", "Infrastructure Requirements")

add_body(
    "All services are recommended to be provisioned on Microsoft Azure to leverage native integrations, "
    "enterprise SLAs, and UNIDO's existing cloud agreements."
)

doc.add_paragraph()

add_info_card("AZU", "0066B3", "Azure App Service",
    "Hosts the Node.js Express backend. Auto-scaling enabled (B2 or higher). Deployment slots for zero-downtime releases.",
    ["Node.js 20 LTS runtime", "Auto-scale: 2–10 instances", "Deployment slots (staging/prod)", "Custom domain + TLS"])

add_info_card("MON", "059669", "MongoDB Atlas",
    "Fully managed document database for sessions, logs, admin data, and conversation history.",
    ["Dedicated M10+ cluster", "Automatic backups (daily)", "VNet peering with Azure", "Encryption at rest"],
    bullet_color=GREEN)

add_info_card("AZU", "33537E", "Azure OpenAI Service",
    "Provides GPT-4o-mini for response generation and text-embedding-3-small for semantic embeddings.",
    ["GPT-4o-mini deployment", "text-embedding-3-small", "Private endpoint access", "Quota: 120k TPM recommended"])

add_info_card("ELA", "E88C30", "Elasticsearch (Elastic Cloud)",
    "Hosts the kNN vector index (1,536 dimensions) for semantic search over UNIDO career content.",
    ["Elastic Cloud (Azure region)", "8.x with kNN plugin", "Hot-warm architecture", "Index lifecycle management"],
    bullet_color=ORANGE)

add_info_card("AZU", "E88C30", "Azure Blob Storage",
    "Durable, geo-redundant storage for raw scraped JSON backups and pipeline artefacts.",
    ["LRS or GRS replication", "Lifecycle policies (90-day retention)", "SAS token access", "Soft delete enabled"],
    bullet_color=ORANGE)

add_info_card("AZU", "C0392B", "Azure Key Vault",
    "Centralised secret management. All API keys, connection strings, and JWT secrets stored and rotated here.",
    ["All app secrets managed here", "Managed identity access", "Audit logging for all reads", "Automatic secret rotation"],
    bullet_color=RED_DARK)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 12 — DEPLOYMENT STRATEGY
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("7", "Deployment Strategy")

add_body(
    "The system supports a staged rollout strategy to minimise risk and validate functionality "
    "at each milestone before progressing to full production."
)

doc.add_paragraph()

add_deployment_phase("Phase 1 — Infrastructure Provisioning", "Week 1-2", "33537E", [
    "Provision Azure App Service, MongoDB Atlas, Elasticsearch",
    "Configure Azure OpenAI deployments",
    "Set up Azure Key Vault with all secrets",
    "Establish network security groups and VNet peering",
])

add_deployment_phase("Phase 2 — Backend Deployment & Testing", "Week 3-4", "0066B3", [
    "Deploy Node.js backend to App Service (staging slot)",
    "Run integration tests (Postman / Jest)",
    "Verify rate limiting, JWT auth, and Socket.IO",
    "Load test at 2× expected peak traffic",
])

add_deployment_phase("Phase 3 — Data Pipeline Run", "Week 5", "059669", [
    "Execute initial full scrape of UNIDO Careers website",
    "Validate chunking, embedding quality, and index count",
    "Test semantic search with domain-specific queries",
    "Review and tune retrieval accuracy (cosine similarity thresholds)",
])

add_deployment_phase("Phase 4 — Frontend Integration", "Week 6", "E88C30", [
    "Deploy React frontend with chatbot widget embedded",
    "End-to-end smoke testing on staging environment",
    "UAT with UNIDO stakeholders",
    "Accessibility audit (WCAG 2.1 AA)",
])

add_deployment_phase("Phase 5 — Production Go-Live", "Week 7", "C0392B", [
    "Swap staging → production deployment slot",
    "Enable monitoring alerts (Azure Monitor + App Insights)",
    "Enable Azure WAF / DDoS Standard",
    "Configure automated scraping schedule (default: daily)",
])


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 13 — ANALYTICS & MONITORING
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("8", "Analytics & Monitoring")

add_body(
    "Comprehensive observability is built into every system component. Real-time dashboards allow "
    "administrators to monitor health, usage, and quality simultaneously."
)

doc.add_paragraph()
add_stat_boxes([
    ("500+", "capacity target\nQueries / Day"),
    ("< 2s", "end-to-end\nP95 Latency"),
    ("99.9%", "SLA\nUptime Target"),
    ("< 1%", "guardrail rate\nBlocked Queries"),
])

add_sub_heading("8.1", "What Is Monitored")
add_styled_table(
    ["Metric", "Description", "Target"],
    [
        ["Chat API Response Time", "Per-request latency tracked via APM", "< 2 s P95"],
        ["Embedding Latency", "Time to generate query vector", "< 500 ms"],
        ["Search Query Latency", "Elasticsearch kNN retrieval time", "< 300 ms"],
        ["Session Creation Rate", "New sessions per minute", "Baseline alert"],
        ["Error Rate", "4xx / 5xx ratio per endpoint", "< 0.5%"],
        ["Guard Block Rate", "% queries blocked by content filter", "Alert if > 5%"],
        ["Data Pipeline Duration", "End-to-end scrape + embed + index time", "< 10 min"],
        ["MongoDB Connection Pool", "Active vs available connections", "Alert at 80%"],
    ],
    col_widths=[4.5, 6.5, 3.5]
)

doc.add_paragraph()
add_sub_heading("8.2", "Alerting Strategy")
alert_items = [
    "Azure Monitor action groups send email + Teams notifications on critical alerts",
    "P1 alerts (service down, >10% error rate): PagerDuty / on-call rotation",
    "P2 alerts (high latency, guard block spike): admin dashboard notification",
    "Weekly automated report: top 20 queries, satisfaction rate, new data coverage",
    "Monthly review: embedding quality scores, retrieval precision@K metrics",
]
for a in alert_items:
    add_bullet(a)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 14-15 — WHAT WE NEED FROM UNIDO
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("9", "What We Need From UNIDO to Go Live")

add_body(
    "To successfully deploy the UNIDO RAG Chatbot into production, MouriTech requires the following access, "
    "approvals, and artefacts from UNIDO. These are pre-requisites that must be in place before the go-live date."
)

doc.add_paragraph()

add_info_card("AZU", "0066B3", "Azure OpenAI API",
    "Powers the chatbot's intelligence to understand questions and generate accurate, human-like responses.",
    ["API key & endpoint for GPT-4o / GPT-4o-mini", "Active Azure OpenAI resource provisioned",
     "Sufficient TPM quota for production load", "Model name, version, and region confirmed"])

add_info_card("ELA", "E88C30", "Elasticsearch (Vector DB)",
    "Enables fast semantic search to find the most relevant career information from scraped data.",
    ["Cluster URL and authentication credentials", "Index config aligned to embedding dimensions",
     "Storage & compute quota for indexed content", "Network access from Azure-hosted services"],
    bullet_color=ORANGE)

add_info_card("AZU", "E88C30", "Azure Blob Storage",
    "Stores raw scraped content and backups for processing and future re-indexing.",
    ["Storage account name and connection string", "Container names for scraped data & backups",
     "RBAC roles assigned to application identity", "Data retention & lifecycle policy agreed"],
    bullet_color=ORANGE)

add_info_card("AZU", "E88C30", "Azure AI Content Safety / Guardrails",
    "Ensures responses are safe, compliant, and filtered for inappropriate or sensitive content.",
    ["Content Safety resource provisioned in Azure", "API key and endpoint provided to app team",
     "Severity threshold config reviewed & approved", "Filtering policy sign-off from compliance team"],
    bullet_color=ORANGE)

add_info_card("MON", "059669", "MongoDB (credentials)",
    "Stores structured data such as scraped content, chatbot logs, configurations, and analytics.",
    ["Connection string with read/write roles", "Database and collection naming confirmed",
     "Azure service IPs whitelisted in MongoDB", "Backup schedule and retention agreed"],
    bullet_color=GREEN)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE 16 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_section_heading("10", "Conclusion")

add_body(
    "The UNIDO RAG Chatbot represents a significant step forward in intelligent public sector digital services. "
    "By combining the power of Azure OpenAI, Elasticsearch semantic search, and a robust data pipeline, MouriTech "
    "has delivered a system that is both technically sophisticated and immediately accessible to UNIDO's global audience."
)

doc.add_paragraph()
add_sub_heading("", "Strengths Delivered", GREEN)
strengths = [
    "Retrieval-grounded answers: zero hallucinations on verified data",
    "Enterprise security: JWT, RBAC, rate limiting, CORS, content filtering",
    "Self-maintaining data: automated pipeline keeps content current",
    "Scalable on Azure: auto-scaling App Service, managed database and search",
    "Operational visibility: full audit logs, real-time admin dashboard",
    "Future-ready: modular architecture supports new capabilities (multilingual, voice, etc.)",
]
for s in strengths:
    add_bullet(s, color=GREEN)

doc.add_paragraph()
add_sub_heading("", "Remaining Go-Live Items", ORANGE)
remaining = [
    "HTTPS / TLS enforcement via Azure Front Door or App Gateway",
    "Azure WAF (Web Application Firewall) + DDoS Standard plan activation",
    "Private VNet peering between App Service, MongoDB, and Elasticsearch",
    "Formal data retention and deletion policy documentation",
    "Final UAT sign-off from UNIDO stakeholders",
]
for r in remaining:
    add_bullet(r, color=ORANGE)

doc.add_paragraph()
doc.add_paragraph()

# System Readiness callout
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
cell = tbl.cell(0, 0)
cell.text = ""
set_cell_shading(cell, "0D2B52")
p = cell.paragraphs[0]
run = p.add_run("  System Readiness Assessment")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = BLUE_MID
p2 = cell.add_paragraph()
run2 = p2.add_run(
    "  Once the remaining hardening items are completed, the UNIDO RAG Chatbot will meet organisational-level "
    "security and operational readiness for live deployment. MouriTech is committed to supporting UNIDO through "
    "go-live and beyond."
)
run2.font.size = Pt(10)
run2.font.color.rgb = WHITE

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by MouriTech Global Enterprise Solutions  |  Version 1.0  |  April 2026  |  Confidential")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = GREY_MID


# ── SAVE ──
output_path = os.path.join(os.path.dirname(__file__), "..", "docs", "UNIDO_RAG_Chatbot_Technical_Documentation.docx")
doc.save(output_path)
print(f"Word document saved to: {os.path.abspath(output_path)}")
