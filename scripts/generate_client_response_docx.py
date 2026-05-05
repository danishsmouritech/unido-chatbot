"""
Generate a professional Word document answering UNIDO's client questions
about hosting, infrastructure, CI/CD, costs, operations, security, and data pipeline.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colours ──
BLUE_DARK  = RGBColor(0x0A, 0x1F, 0x3D)
BLUE_MID   = RGBColor(0x00, 0x66, 0xB3)
GREEN      = RGBColor(0x05, 0x96, 0x69)
ORANGE     = RGBColor(0xE8, 0x8C, 0x30)
RED_DARK   = RGBColor(0xC0, 0x39, 0x2B)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_DARK  = RGBColor(0x33, 0x41, 0x55)
GREY_MID   = RGBColor(0x64, 0x74, 0x8B)
BLACK      = RGBColor(0x1E, 0x29, 0x3B)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = GREY_DARK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val["sz"]}" '
                f'w:space="0" w:color="{val["color"]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)

def remove_table_borders(tbl):
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_horizontal_line(color="0066B3"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, color)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pf = p._element
        pPr = pf.get_or_add_pPr()
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="40" w:lineRule="exact"/>')
        pPr.append(spacing)
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="40" w:hRule="exact"/>')
    trPr.append(trHeight)
    remove_table_borders(tbl)
    return tbl


def add_section_heading(number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run_bar = p.add_run("█  ")
    run_bar.font.color.rgb = BLUE_MID
    run_bar.font.size = Pt(18)
    run = p.add_run(f"{number}. {title}")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = GREY_DARK

def add_sub_heading(title, color=BLUE_MID):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = color

def add_sub_sub_heading(title, color=GREY_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = color

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = GREY_DARK
    return p

def add_bullet(text, bold_prefix=None, color=BLUE_MID):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0)
    run_dot = p.add_run("●  ")
    run_dot.font.size = Pt(7)
    run_dot.font.color.rgb = color
    if bold_prefix:
        run_b = p.add_run(bold_prefix + " ")
        run_b.font.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.color.rgb = GREY_DARK
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = GREY_DARK
    return p

def add_callout(text, accent_color="0066B3", bg_color="EBF5FF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)
    cell = tbl.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, bg_color)
    set_cell_borders(cell, left={"sz": "16", "color": accent_color})
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GREY_DARK
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

def add_styled_table(headers, rows, col_widths=None, header_color="33537E"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, header_color)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = GREY_DARK
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("MOURITECH  |  GLOBAL ENTERPRISE SOLUTIONS")
run.font.bold = True
run.font.size = Pt(11)
run.font.color.rgb = BLUE_MID
add_horizontal_line("0066B3")
add_horizontal_line("E88C30")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("UNIDO AI Chatbot\nClient Queries — Detailed Response Document")
run.font.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE_DARK

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    "Comprehensive answers to UNIDO's questions on hosting & infrastructure, system landscape, "
    "CI/CD, cost structure, operations, security compliance, and data pipeline maintenance — "
    "all grounded in the actual system implementation."
)
run.font.size = Pt(12)
run.font.color.rgb = GREY_DARK

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Version", "1.0"),
    ("Date", "April 2026"),
    ("Classification", "Internal — Organizational"),
    ("Prepared by", "MouriTech Global Enterprise Solutions"),
    ("Client", "UNIDO — United Nations Industrial Development Organization"),
]
tbl = doc.add_table(rows=len(meta), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
for i, (key, val) in enumerate(meta):
    for ci in range(3):
        tbl.cell(i, ci).text = ""
    p_k = tbl.cell(i, 0).paragraphs[0]
    run_k = p_k.add_run(key)
    run_k.font.size = Pt(10); run_k.font.color.rgb = GREY_MID
    p_sep = tbl.cell(i, 1).paragraphs[0]
    run_sep = p_sep.add_run(":")
    run_sep.font.size = Pt(10); run_sep.font.color.rgb = GREY_MID
    p_v = tbl.cell(i, 2).paragraphs[0]
    run_v = p_v.add_run(val)
    run_v.font.bold = True; run_v.font.size = Pt(10); run_v.font.color.rgb = GREY_DARK
    if i < len(meta) - 1:
        for ci in range(3):
            set_cell_borders(tbl.cell(i, ci), bottom={"sz": "2", "color": "E2E8F0"})

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL  |  MouriTech  |  UNIDO AI Chatbot  |  Version 1.0  |  April 2026")
run.font.size = Pt(8); run.font.color.rgb = GREY_MID; run.font.bold = True


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Table of Contents")
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = GREY_DARK
add_horizontal_line("0066B3")
doc.add_paragraph()

toc = [
    ("1.", "Hosting & Infrastructure Setup", True),
    ("  1.1", "Recommended Hosting Option", False),
    ("  1.2", "Minimum & Recommended Specifications", False),
    ("  1.3", "Runtime & OS Requirements", False),
    ("  1.4", "Rationale", False),
    ("2.", "System Landscape & Environments", True),
    ("  2.1", "Environment Strategy", False),
    ("  2.2", "Component Interaction & Data Flow", False),
    ("3.", "CI/CD & Deployment Process", True),
    ("  3.1", "Application CI/CD Pipeline", False),
    ("  3.2", "Infrastructure as Code (IaC)", False),
    ("4.", "Cost Structure", True),
    ("  4.1", "Fixed Infrastructure Costs", False),
    ("  4.2", "Variable / Consumption-Based Costs", False),
    ("  4.3", "Monthly Cost Estimates", False),
    ("  4.4", "Cost Optimisation Mechanisms", False),
    ("5.", "Operations & Support Model", True),
    ("  5.1", "Incident Management Procedures", False),
    ("  5.2", "Backup, Recovery & Disaster Recovery", False),
    ("6.", "Security & Compliance (Operational View)", True),
    ("  6.1", "Data Retention & Deletion Policies", False),
    ("7.", "Data Pipeline & Maintenance", True),
    ("  7.1", "Scraping Scope & Dependencies", False),
    ("  7.2", "Re-indexing Strategy & Operational Triggers", False),
    ("  7.3", "Manual vs. Automated Maintenance Tasks", False),
]

tbl = doc.add_table(rows=len(toc), cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
for i, (num, title, is_main) in enumerate(toc):
    cell_num = tbl.cell(i, 0); cell_num.width = Cm(1.5); cell_num.text = ""
    p = cell_num.paragraphs[0]
    run = p.add_run(num)
    run.font.bold = is_main; run.font.size = Pt(10.5) if is_main else Pt(10)
    run.font.color.rgb = GREY_DARK if is_main else GREY_MID
    cell_title = tbl.cell(i, 1); cell_title.text = ""
    p2 = cell_title.paragraphs[0]
    run2 = p2.add_run(title)
    run2.font.bold = is_main; run2.font.size = Pt(10.5) if is_main else Pt(10)
    run2.font.color.rgb = GREY_DARK if is_main else GREY_MID
    if is_main:
        set_cell_borders(cell_num, bottom={"sz": "2", "color": "E2E8F0"})
        set_cell_borders(cell_title, bottom={"sz": "2", "color": "E2E8F0"})


# ══════════════════════════════════════════════════════════════════════════════
# 1. HOSTING & INFRASTRUCTURE SETUP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("1", "Hosting & Infrastructure Setup")

add_body(
    "This section specifies the exact server and hosting requirements derived directly from the application's "
    "codebase, dependency footprint, and runtime behaviour. Recommendations are production-grade and aligned "
    "with Azure's enterprise hosting capabilities."
)

# 1.1
add_sub_heading("1.1  Recommended Hosting Option")

add_body(
    "Azure App Service (Linux, B2 or P1v3 tier) is the recommended hosting platform for the UNIDO RAG Chatbot backend."
)

add_styled_table(
    ["Hosting Option", "Verdict", "Rationale"],
    [
        ["Azure App Service (Linux)", "✅ Recommended", "Managed PaaS, auto-scale, deployment slots, built-in TLS, native Azure OpenAI integration, zero infrastructure overhead"],
        ["Azure Container App", "✅ Good Alternative", "Suitable if containerisation is preferred; supports scale-to-zero for cost savings on non-production environments"],
        ["Azure VM (IaaS)", "⚠️ Not Recommended", "Requires manual OS patching, scaling, and monitoring. Higher operational burden with no architectural benefit for this workload"],
        ["Azure Kubernetes Service (AKS)", "❌ Over-engineered", "Designed for multi-service microservices at scale. This is a single-service application — AKS adds complexity and cost without proportional benefit"],
    ],
    col_widths=[4, 3, 9.5]
)

add_callout(
    "Rationale: The chatbot is a single Node.js Express application with a Puppeteer-based scraping pipeline. "
    "Azure App Service provides the simplest path to production with built-in TLS, deployment slots for zero-downtime "
    "releases, auto-scaling, and native integration with Azure Monitor. Puppeteer requires a Linux host with Chromium "
    "dependencies — App Service Linux satisfies this without custom Docker images."
)

# 1.2
add_sub_heading("1.2  Minimum & Recommended Specifications")

add_sub_sub_heading("Backend Server (Azure App Service)")
add_styled_table(
    ["Resource", "Minimum (Dev/Staging)", "Recommended (Production)", "Reason"],
    [
        ["App Service Plan", "B2 (Basic)", "P1v3 (Premium v3)", "P1v3 provides 2 vCPU, 8 GB RAM, deployment slots, auto-scale, and VNet integration"],
        ["vCPU", "2 cores", "2–4 cores", "Puppeteer (headless Chromium) is CPU-intensive during daily scraping. Normal chat operations are lightweight (~0.2 vCPU)"],
        ["RAM", "3.5 GB", "8 GB", "Puppeteer/Chromium consumes ~500 MB–1 GB during scraping. Node.js Express + Socket.IO needs ~200–400 MB at steady state. 8 GB provides headroom for concurrent operations"],
        ["Storage", "10 GB", "30 GB", "Application code (~5 MB), node_modules + Puppeteer/Chromium (~450 MB), local JSON backups (~50–100 MB rotating), logs. 30 GB allows growth"],
        ["OS", "Linux (Ubuntu 22.04+)", "Linux (Ubuntu 22.04+)", "Puppeteer requires Linux for headless Chromium. Express 5 requires Node.js ≥18"],
        ["Runtime", "Node.js 20 LTS", "Node.js 20 LTS", "Express 5.2.1 requires Node.js ≥18. LTS 20 is the current supported version with security patches until April 2026"],
        ["Auto-scale", "1 instance", "2–5 instances", "Scale on CPU >70% threshold. Minimum 2 instances for high availability in production"],
    ],
    col_widths=[3, 3.5, 3.5, 6.5]
)

add_sub_sub_heading("Frontend (Static Site)")
add_styled_table(
    ["Resource", "Specification", "Reason"],
    [
        ["Hosting", "Azure Static Web Apps (Free/Standard) or Azure Blob Storage + CDN", "React 19 + Vite builds to static HTML/JS/CSS. No server-side rendering required"],
        ["Storage", "< 50 MB", "Vite production build output is minimal (~2–5 MB). Asset caching via CDN"],
        ["CDN", "Azure Front Door or Azure CDN", "Global distribution for UNIDO's worldwide candidate audience. Also provides WAF and DDoS protection"],
    ],
    col_widths=[3, 5, 8.5]
)

add_sub_sub_heading("Managed Services")
add_styled_table(
    ["Service", "Specification", "Reason"],
    [
        ["MongoDB Atlas", "M10 dedicated cluster (2 GB RAM, 10 GB storage)", "4 collections: chatSessions, chatLogs, adminUsers, adminSettings. Low-moderate write volume. M10 provides automatic backups and encryption at rest"],
        ["Elasticsearch", "Elastic Cloud — 4 GB RAM, 2 vCPU (1 node)", "Single index (unido_careers_index) with ~2,000–5,000 dense_vector documents (1536 dimensions). kNN search requires RAM for vector index in memory"],
        ["Azure OpenAI", "GPT-4o-mini + text-embedding-3-small deployments", "Both models deployed in same Azure region as App Service. Recommended quota: 120K TPM for embeddings, 60K TPM for chat"],
        ["Azure Blob Storage", "LRS (Standard), 1 container", "Stores date-prefixed JSON backups of scraped data and chunks. Current retention: 7 days auto-cleanup. Minimal storage (~100 MB)"],
        ["Azure Key Vault", "Standard tier", "Stores MONGO_URI, AZURE_OPENAI_KEY, JWT_SECRET, ELASTIC credentials. Managed identity access from App Service"],
    ],
    col_widths=[3, 5, 8.5]
)

# 1.3
add_sub_heading("1.3  Runtime & OS Requirements")
add_styled_table(
    ["Component", "Requirement", "Details"],
    [
        ["Node.js", "≥ 18.x (recommended: 20 LTS)", "Express 5.2.1 mandates Node.js ≥18. Node 20 LTS has security support until April 2026"],
        ["NPM", "≥ 9.x", "Ships with Node.js 20 LTS. Required to install 24 production dependencies"],
        ["Chromium", "Bundled with Puppeteer 24.x", "Puppeteer downloads its own Chromium (~170 MB compressed, ~450 MB installed). No system Chrome needed"],
        ["Linux Packages", "libnss3, libatk1.0, libx11, libgbm", "Required by Chromium on Linux. Azure App Service Linux images include these by default"],
        ["TLS/SSL", "TLS 1.2+", "Enforced via HSTS (max-age: 1 year, includeSubDomains, preload) and Helmet middleware"],
    ],
    col_widths=[3, 4.5, 9]
)

# 1.4
add_sub_heading("1.4  Rationale")
add_body("The specification is driven by three workload profiles the application exhibits:")

add_bullet("Steady-state chat serving (95% of time): lightweight — Express handles HTTP requests, Socket.IO manages WebSockets, MongoDB reads/writes are small. CPU < 20%, RAM < 500 MB.", bold_prefix="Normal Operations:", color=GREEN)
add_bullet("Daily pipeline execution (1× per day, ~5–10 minutes): CPU-intensive — Puppeteer launches headless Chromium to scrape 6 data sources, generates ~2,000+ embeddings via Azure OpenAI API, and bulk-indexes into Elasticsearch. Peak CPU ~80%, RAM ~1.5–2 GB.", bold_prefix="Pipeline Scraping:", color=ORANGE)
add_bullet("Admin dashboard with real-time Socket.IO updates, CSV exports, and MongoDB aggregation queries. Lightweight but concurrent with chat traffic.", bold_prefix="Admin Operations:", color=BLUE_MID)

add_callout(
    "The P1v3 App Service tier (2 vCPU, 8 GB RAM) comfortably handles all three workload profiles concurrently. "
    "For cost-conscious deployments, a B2 tier (2 vCPU, 3.5 GB RAM) works for dev/staging environments where the "
    "pipeline runs less frequently.",
    "059669", "ECFDF5"
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM LANDSCAPE & ENVIRONMENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("2", "System Landscape & Environments")

add_sub_heading("2.1  Environment Strategy")
add_body(
    "We recommend a three-environment setup (Development, Staging, Production) leveraging Azure App Service deployment "
    "slots and environment-specific configuration via Azure Key Vault and App Settings."
)

add_styled_table(
    ["Environment", "Purpose", "Infrastructure", "Data", "Access"],
    [
        ["Development", "Active development & testing", "App Service B1 (1 instance), MongoDB Atlas M0 (free), Local Elasticsearch", "Synthetic/sample data, no real user data", "Dev team only, no public access"],
        ["Staging", "Pre-production validation, UAT", "App Service B2 (1 instance), MongoDB Atlas M10, Elastic Cloud (dev tier)", "Clone of production structure, anonymised data", "Dev team + UNIDO UAT testers, IP-restricted"],
        ["Production", "Live system serving UNIDO users", "App Service P1v3 (2–5 instances, auto-scale), MongoDB Atlas M10+, Elastic Cloud (production)", "Real scraped UNIDO data, live user sessions", "Public (chat widget), Admin (JWT + IP whitelist)"],
    ],
    col_widths=[2.5, 3.5, 4.5, 3.5, 3]
)

add_callout(
    "Staging uses Azure App Service deployment slots — this enables zero-downtime deployments by swapping staging "
    "↔ production slots after validation. Each environment has isolated MongoDB databases and Elasticsearch indices "
    "to prevent data cross-contamination.",
    "0066B3", "EBF5FF"
)

add_sub_heading("2.2  Component Interaction & Data Flow")
add_body("The system landscape consists of six interconnected components. All communication occurs over TLS-encrypted channels.")

add_sub_sub_heading("External-Facing Components")
add_bullet("React 19 Frontend (Azure Static Web Apps / CDN) — serves the chat widget and admin dashboard to browsers", bold_prefix="Frontend SPA:")
add_bullet("Node.js Express 5 Backend (Azure App Service) — REST API + Socket.IO server, enforces all security middleware", bold_prefix="Backend API:")

add_sub_sub_heading("Internal Services")
add_bullet("MongoDB Atlas — stores chat sessions, conversation logs, admin users, admin settings (4 collections)", bold_prefix="MongoDB:")
add_bullet("Elastic Cloud — single index (unido_careers_index) with 1536-dim dense_vector fields for kNN cosine search", bold_prefix="Elasticsearch:")
add_bullet("GPT-4o-mini (chat completion, temp 0.1) + text-embedding-3-small (1536-dim embeddings)", bold_prefix="Azure OpenAI:")
add_bullet("LRS storage for date-prefixed JSON backups of scraped data (7-day retention)", bold_prefix="Azure Blob Storage:")

add_sub_sub_heading("Data Flow: Chat Query (Real-Time)")
add_body(
    "Browser → [HTTPS] → Azure CDN/Front Door → [HTTPS] → App Service (Express) → Guard Service → "
    "Azure OpenAI (embed) → Elasticsearch (kNN search) → Azure OpenAI (generate) → MongoDB (log) → "
    "Express → [WebSocket] → Browser"
)

add_sub_sub_heading("Data Flow: Pipeline (Daily Midnight UTC)")
add_body(
    "Node-Cron trigger → Puppeteer (scrapes careers.unido.org + unido.org/employment/faqs) → Cheerio (clean HTML) → "
    "Chunker (1000-char splits, 1-sentence overlap) → Azure OpenAI (embed each chunk) → Elasticsearch (bulk index, "
    "full re-index) → Azure Blob Storage (JSON backup) → MongoDB (update lastScrapeAt)"
)

add_sub_sub_heading("Data Flow: Admin Operations")
add_body(
    "Admin Browser → [HTTPS + JWT Bearer] → App Service → MongoDB (analytics aggregation, settings CRUD, "
    "CSV export) → Socket.IO broadcast (analytics:updated, chatbot:visibilityChanged) → All connected clients"
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. CI/CD & DEPLOYMENT PROCESS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("3", "CI/CD & Deployment Process")

add_sub_heading("3.1  Application CI/CD Pipeline")
add_body(
    "We recommend GitHub Actions (or Azure DevOps Pipelines) for continuous integration and deployment. "
    "The pipeline follows a trunk-based development model with branch protection on main."
)

add_sub_sub_heading("Backend Pipeline (Node.js)")
add_styled_table(
    ["Stage", "Trigger", "Actions", "Environment"],
    [
        ["Build", "Push to feature/* or PR to main", "npm ci → npm run lint → npm audit --audit-level=high", "CI runner"],
        ["Test", "Push to feature/* or PR to main", "npm test (Jest unit + integration tests) → coverage report", "CI runner"],
        ["Deploy to Staging", "Merge to main", "npm ci --production → zip artifact → deploy to App Service staging slot", "Staging"],
        ["Smoke Test", "Post staging deploy", "Health check (GET /health) → API integration tests against staging", "Staging"],
        ["Production Release", "Manual approval gate", "Swap staging ↔ production deployment slot (zero-downtime)", "Production"],
        ["Rollback", "On failure", "Swap slots back to previous version (< 30 seconds)", "Production"],
    ],
    col_widths=[3, 3.5, 6, 3.5]
)

add_sub_sub_heading("Frontend Pipeline (React + Vite)")
add_styled_table(
    ["Stage", "Trigger", "Actions", "Environment"],
    [
        ["Build", "Push to feature/* or PR to main", "npm ci → npm run build (Vite production build) → output to dist/", "CI runner"],
        ["Deploy to Staging", "Merge to main", "Upload dist/ to Azure Static Web Apps (staging environment)", "Staging"],
        ["Production Release", "Manual approval or auto-promote", "Promote staging → production in Azure Static Web Apps", "Production"],
    ],
    col_widths=[3, 3.5, 6, 3.5]
)

add_callout(
    "Zero-Downtime Deployment: Azure App Service deployment slots allow the staging slot to be warmed up and "
    "validated before swapping to production. The swap operation is atomic and completes in < 30 seconds. "
    "If issues are detected post-swap, an immediate rollback is available by swapping again.",
    "059669", "ECFDF5"
)

add_sub_heading("3.2  Infrastructure as Code (IaC)")
add_body(
    "All Azure infrastructure should be provisioned and managed via Terraform or Azure Bicep templates "
    "stored in the same Git repository under an infra/ directory."
)

add_styled_table(
    ["Component", "IaC Resource", "Configuration Managed"],
    [
        ["App Service Plan", "azurerm_service_plan", "SKU (P1v3), OS (Linux), auto-scale rules (CPU > 70%, min 2, max 5)"],
        ["App Service", "azurerm_linux_web_app", "Node 20 LTS runtime, deployment slots, app settings, VNet integration"],
        ["Static Web App", "azurerm_static_web_app", "Frontend hosting, custom domain, API integration"],
        ["Key Vault", "azurerm_key_vault", "Secrets for MONGO_URI, AZURE_OPENAI_KEY, JWT_SECRET, ES credentials"],
        ["Blob Storage", "azurerm_storage_account + container", "LRS replication, lifecycle policy (7-day delete), SAS tokens"],
        ["Azure Front Door", "azurerm_cdn_frontdoor_*", "CDN, WAF policy, DDoS protection, custom domain + TLS"],
        ["Monitor / Alerts", "azurerm_monitor_*", "Action groups, metric alerts (CPU, memory, error rate, response time)"],
    ],
    col_widths=[3, 4, 9.5]
)

add_callout(
    "Recommendation: Store Terraform state in Azure Blob Storage with state locking via Azure Table Storage. "
    "Run terraform plan in PR checks and terraform apply only on merge to the infrastructure branch.",
    "E88C30", "FFFDF5"
)


# ══════════════════════════════════════════════════════════════════════════════
# 4. COST STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("4", "Cost Structure")

add_body(
    "All costs below are based on Azure East US / West Europe pricing (April 2026). Actual costs may vary "
    "based on UNIDO's Azure Enterprise Agreement discounts and reserved instance commitments."
)

add_sub_heading("4.1  Fixed Infrastructure Costs (Monthly)")
add_styled_table(
    ["Service", "Tier / Config", "Monthly Cost (USD)", "Notes"],
    [
        ["Azure App Service", "P1v3 Linux (2 vCPU, 8 GB, 250 GB storage)", "$140–160", "Single instance. Add ~$140/instance for auto-scale replicas"],
        ["MongoDB Atlas", "M10 Dedicated (2 GB RAM, 10 GB, AWS/Azure)", "$55–80", "Includes automated backups. M0 free tier for dev/staging"],
        ["Elasticsearch (Elastic Cloud)", "4 GB RAM, 2 vCPU, 120 GB storage", "$95–120", "Single node. Sufficient for ~5,000 vector documents (1536-dim)"],
        ["Azure Blob Storage", "LRS Standard, ~1 GB", "$2–5", "Minimal: JSON backups rotated every 7 days"],
        ["Azure Key Vault", "Standard, ~100 operations/day", "$0.03/10K ops ≈ $1", "Secrets retrieval at app startup + token refresh"],
        ["Azure Static Web Apps", "Standard tier", "$9", "Frontend hosting with custom domain + CDN"],
        ["Azure Front Door (CDN + WAF)", "Standard tier", "$35–50", "Global CDN + managed WAF rules + DDoS protection"],
        ["Azure Monitor + App Insights", "5 GB/month data ingestion", "$0 (included) – $30", "First 5 GB/month free. Alerts included"],
        ["", "TOTAL (Fixed)", "$337–455 / month", "Production environment estimate"],
    ],
    col_widths=[3.5, 4.5, 3, 5.5]
)

add_sub_heading("4.2  Variable / Consumption-Based Costs (Monthly)")

add_sub_sub_heading("Azure OpenAI Token Costs")
add_body(
    "Token usage depends on chat volume and pipeline frequency. Below are estimates based on the actual "
    "application's token consumption patterns:"
)

add_styled_table(
    ["Operation", "Model", "Tokens Per Operation", "Daily Volume", "Monthly Tokens", "Cost / 1M Tokens", "Monthly Cost (USD)"],
    [
        ["Chat — Embedding", "text-embedding-3-small", "~200 tokens/query", "100–300 queries", "~2M tokens", "$0.02", "$0.04"],
        ["Chat — Generation (input)", "GPT-4o-mini (input)", "~2,000 tokens (prompt + context + history)", "100–300 queries", "~18M tokens", "$0.15", "$2.70"],
        ["Chat — Generation (output)", "GPT-4o-mini (output)", "~300 tokens/response", "100–300 queries", "~2.7M tokens", "$0.60", "$1.62"],
        ["Pipeline — Embed chunks", "text-embedding-3-small", "~500 tokens/chunk × 2,000–5,000 chunks", "1× daily", "~75M tokens", "$0.02", "$1.50"],
        ["", "", "", "", "TOTAL (Variable)", "", "$5.86 / month"],
    ],
    col_widths=[3, 2.5, 3, 2, 2, 2, 2]
)

add_callout(
    "Azure OpenAI costs are remarkably low for this workload. GPT-4o-mini is one of the most cost-efficient "
    "models available. Even at 3× current estimated volume, total AI costs remain under $20/month. "
    "The system's guard service also reduces costs by blocking ~1–5% of queries before they reach the AI layer.",
    "059669", "ECFDF5"
)

add_sub_heading("4.3  Monthly Cost Estimates — Scenario Summary")
add_styled_table(
    ["Scenario", "Chat Volume", "Infrastructure", "Azure OpenAI", "Total Monthly"],
    [
        ["Development", "Minimal testing", "~$60 (B1 + free Atlas + local ES)", "< $1", "$60–65"],
        ["Staging", "UAT testing only", "~$150 (B2 + M10 + Elastic dev)", "< $2", "$150–155"],
        ["Production — Expected Load", "100–300 queries/day", "$337–455", "$5–10", "$345–465"],
        ["Production — Peak Load", "500–1,000 queries/day", "$480–600 (2 App Service instances)", "$15–30", "$500–630"],
        ["Production — Scale (future)", "2,000+ queries/day", "$700–900 (3–5 instances + M20 Atlas)", "$40–80", "$740–980"],
    ],
    col_widths=[3.5, 3.5, 4, 2.5, 3]
)

add_sub_heading("4.4  Cost Optimisation Mechanisms")
add_bullet("Azure Reserved Instances for App Service (1-year commitment) — saves 30–40% on compute costs", bold_prefix="Reserved Pricing:", color=GREEN)
add_bullet("MongoDB Atlas M10 free tier for dev/staging. Production M10 is $57/month — smaller than most database costs", bold_prefix="Right-Sized Database:", color=GREEN)
add_bullet("Guard service blocks offensive, confidential, prompt injection, and out-of-scope queries before they hit Azure OpenAI — eliminating wasted tokens", bold_prefix="Guard Service Savings:", color=GREEN)
add_bullet("GPT-4o-mini is 15× cheaper than GPT-4o with comparable quality for grounded RAG responses", bold_prefix="Model Selection:", color=GREEN)
add_bullet("Only chunks scoring ≥ 0.75 cosine similarity are sent to GPT, reducing average context window size and token usage", bold_prefix="RAG Min-Score Filter:", color=GREEN)
add_bullet("The entire vector index is rebuilt daily rather than appending, preventing index bloat and maintaining search performance", bold_prefix="Full Re-index Strategy:", color=GREEN)
add_bullet("Auto-scale rules (CPU > 70%) add instances only when needed and scale down during low-traffic periods", bold_prefix="Auto-Scaling:", color=GREEN)
add_bullet("Azure Container Apps (scale-to-zero) can reduce dev/staging costs to near-zero when idle", bold_prefix="Scale-to-Zero (Dev/Staging):", color=GREEN)


# ══════════════════════════════════════════════════════════════════════════════
# 5. OPERATIONS & SUPPORT MODEL
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("5", "Operations & Support Model")

add_sub_heading("5.1  Incident Management Procedures")

add_sub_sub_heading("Severity Classification")
add_styled_table(
    ["Severity", "Definition", "Response Time", "Resolution Target", "Example"],
    [
        ["P1 — Critical", "Service completely down or data breach", "< 15 minutes", "< 2 hours", "App Service crash, MongoDB connection failure, security breach"],
        ["P2 — High", "Significant degradation affecting users", "< 30 minutes", "< 4 hours", "Azure OpenAI quota exceeded, Elasticsearch index corruption, chatbot returning errors > 10%"],
        ["P3 — Medium", "Feature degraded but service operational", "< 2 hours", "< 1 business day", "Pipeline scraping failure, slow response times (> 5s P95), admin dashboard errors"],
        ["P4 — Low", "Minor issue, cosmetic or edge-case", "< 1 business day", "Next sprint", "UI rendering glitch, non-critical log noise, documentation update"],
    ],
    col_widths=[2.5, 3.5, 2.5, 2.5, 5.5]
)

add_sub_sub_heading("Detection & Alerting")
add_bullet("Azure Monitor metric alerts on: CPU > 80%, memory > 75%, HTTP 5xx > 5%, response time P95 > 3s", bold_prefix="Infrastructure Alerts:")
add_bullet("Application Insights tracks: request rate, dependency calls (MongoDB, ES, OpenAI), exception rate, custom events", bold_prefix="Application Monitoring:")
add_bullet("Express audit logger flags: all requests with status ≥ 400, all requests with duration > 5,000ms, failed JWT verifications (with IP logged)", bold_prefix="Built-in Audit Logger:")
add_bullet("MongoDB ChatLog collection stores: every Q&A with status (success/fallback/error), timing breakdown (embedding/retrieval/generation), guard reason for blocked queries", bold_prefix="Chat-Level Monitoring:")

add_sub_sub_heading("Escalation Path")
add_body(
    "Azure Monitor action groups → email + Microsoft Teams notification → L1 (MouriTech operations) → "
    "L2 (MouriTech engineering) → L3 (Azure support for platform issues / Elastic support for search issues)"
)

add_sub_heading("5.2  Backup, Recovery & Disaster Recovery")

add_sub_sub_heading("Backup Strategy")
add_styled_table(
    ["Component", "Backup Method", "Frequency", "Retention", "Recovery Point Objective (RPO)"],
    [
        ["MongoDB Atlas", "Automated continuous snapshots", "Continuous (point-in-time)", "7 days (M10+)", "< 1 minute (any point in last 7 days)"],
        ["Elasticsearch Index", "Re-index from pipeline (source of truth is UNIDO website)", "Rebuilt daily at midnight UTC", "Current + 1 previous (JSON backups)", "< 24 hours (last successful pipeline run)"],
        ["Scraped Data (JSON)", "Local files + Azure Blob Storage", "Every pipeline run", "7 days auto-delete in Blob Storage", "< 24 hours"],
        ["Application Code", "Git repository (GitHub/Azure DevOps)", "Every commit", "Indefinite", "Instant (any commit)"],
        ["Azure OpenAI Config", "Infrastructure as Code (Terraform)", "Every infrastructure change", "Indefinite", "Instant (redeploy)"],
        ["Admin Settings", "Part of MongoDB backup", "Continuous", "7 days", "< 1 minute"],
    ],
    col_widths=[3, 3.5, 3, 3, 4]
)

add_sub_sub_heading("Disaster Recovery")
add_styled_table(
    ["Scenario", "Recovery Procedure", "RTO", "RPO"],
    [
        ["App Service failure", "Auto-restart (built-in). If persistent: redeploy from latest Git commit to new instance", "< 5 min (auto) / < 30 min (manual)", "Zero (stateless app)"],
        ["MongoDB failure", "Atlas automatic failover to replica. Point-in-time restore if data corruption", "< 1 min (failover) / < 15 min (restore)", "< 1 min"],
        ["Elasticsearch index corruption", "Re-run pipeline: node pipeline.js re-scrapes, re-embeds, and re-indexes all data", "< 15 min", "< 24 hours"],
        ["Azure OpenAI outage", "Return cached fallback message: 'Service temporarily unavailable'. Resume when API recovers", "Automatic (graceful degradation)", "N/A"],
        ["Full region outage", "Redeploy to secondary Azure region from IaC + restore MongoDB from cross-region backup", "< 2 hours", "< 1 min (MongoDB), < 24h (ES)"],
    ],
    col_widths=[3, 7.5, 3, 3]
)

add_callout(
    "Key architectural advantage: Elasticsearch is NOT the source of truth — the UNIDO website is. If the ES index "
    "is lost or corrupted, a single pipeline run (~5–10 minutes) completely rebuilds it from scratch by re-scraping "
    "and re-embedding. This eliminates the need for complex ES backup/restore procedures.",
    "059669", "ECFDF5"
)


# ══════════════════════════════════════════════════════════════════════════════
# 6. SECURITY & COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("6", "Security & Compliance (Operational View)")

add_sub_heading("6.1  Data Retention & Deletion Policies")

add_body(
    "The following retention policies are recommended based on the data types stored in the system. "
    "These should be reviewed and formally approved by UNIDO's data governance team before go-live."
)

add_styled_table(
    ["Data Category", "Storage Location", "Current Behaviour", "Recommended Retention", "Deletion Method", "Rationale"],
    [
        ["Chat Sessions", "MongoDB (chatSessions)", "Retained indefinitely", "90 days, then auto-delete", "MongoDB TTL index on createdAt", "Sessions are anonymous (UUID-based). 90 days covers audit needs while limiting data footprint"],
        ["Chat Logs (Audit)", "MongoDB (chatLogs)", "Retained indefinitely", "12 months, then archive/delete", "Scheduled job or MongoDB TTL", "Audit logs are valuable for analytics, security review, and compliance. 12 months aligns with typical audit cycles"],
        ["Admin Users", "MongoDB (adminUsers)", "Retained while active", "Delete when user is deactivated + 30 days", "Manual admin action", "GDPR-aligned: data kept only as long as needed. Deactivated accounts kept 30 days for audit trail"],
        ["Admin Settings", "MongoDB (adminSettings)", "Singleton document, updated in place", "No deletion needed", "N/A", "Single config record; updated, never multiplied"],
        ["Scraped Data (JSON)", "Azure Blob Storage", "7-day auto-delete (when enabled)", "7 days", "Blob lifecycle policy", "Scraped data is public UNIDO website content. Short retention is sufficient as pipeline rebuilds daily"],
        ["Elasticsearch Index", "Elastic Cloud", "Full re-index daily (delete all + rebuild)", "Current index only", "deleteByQuery on each pipeline run", "Index is ephemeral — rebuilt from source (UNIDO website) every 24 hours"],
        ["Request Audit Logs", "Application logs (stdout)", "Streamed to Azure Monitor", "30 days in Log Analytics", "Azure Monitor retention policy", "Standard operational log retention. Can be extended for compliance needs"],
        ["IP Addresses", "MongoDB (chatLogs.requestMeta.ip)", "Retained with chat logs", "Same as chat logs (12 months)", "Deleted with parent log", "Used for rate limiting and security audit. Not PII under UNIDO context (no user accounts)"],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 2.5, 4]
)

add_callout(
    "Important: No PII (Personally Identifiable Information) is collected or stored by the chatbot. "
    "Chat sessions are anonymous (UUID-based), no names/emails/passwords are collected from end users. "
    "Only admin accounts (internal UNIDO staff) have personal data (email + hashed password). "
    "This significantly simplifies GDPR/data protection compliance.",
    "059669", "ECFDF5"
)

add_sub_sub_heading("Implementation Steps Required")
add_bullet("Add MongoDB TTL indexes on chatSessions.createdAt (90 days) and chatLogs.createdAt (365 days)", bold_prefix="Action 1:")
add_bullet("Enable Azure Blob Storage lifecycle policy (already coded, currently commented out — needs AZURE_STORAGE_CONNECTION_STRING)", bold_prefix="Action 2:")
add_bullet("Configure Azure Monitor Log Analytics retention to 30 days (or as per UNIDO's compliance requirements)", bold_prefix="Action 3:")
add_bullet("Document and publish a Data Retention Policy signed off by UNIDO data governance", bold_prefix="Action 4:")


# ══════════════════════════════════════════════════════════════════════════════
# 7. DATA PIPELINE & MAINTENANCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_section_heading("7", "Data Pipeline & Maintenance")

add_sub_heading("7.1  Scraping Scope & Dependencies")

add_sub_sub_heading("Data Sources")
add_styled_table(
    ["Scraper", "Source URL", "Data Extracted", "Dependencies", "Failure Impact"],
    [
        ["Homepage", "https://careers.unido.org/", "Hero title, description, CTA, career category links", "Puppeteer + Cheerio", "Missing homepage info in chatbot responses"],
        ["FAQs", "https://www.unido.org/employment/faqs", "Numbered Q&A pairs", "Puppeteer + Cheerio", "Cannot answer FAQ-type questions"],
        ["Job Search", "https://careers.unido.org/search/", "Job title, location, department, grade, deadline, detail URL. Paginates at 25/page", "Puppeteer (browser navigation)", "No job listing data — high impact"],
        ["Job Details", "Individual job URLs (dynamically discovered)", "Full job description: requisition ID, grade, duty station, responsibilities, qualifications, etc.", "Puppeteer + previous step's URLs", "Incomplete job details — high impact"],
        ["Categories", "Category page URLs (from homepage scrape)", "Category title, description, job count, job listings per category", "Puppeteer + homepage step's URLs", "Missing category browsing data"],
        ["Category Content", "Learn-more page URLs (from categories)", "Detailed paragraphs, lists, tables (grade/experience requirements). 1.5s delay between pages", "Puppeteer + categories step's URLs", "Missing detailed category explanations"],
    ],
    col_widths=[2.5, 3.5, 4, 2.5, 4]
)

add_sub_sub_heading("Scraping Frequency & Schedule")
add_styled_table(
    ["Trigger", "Schedule", "Configuration", "Notes"],
    [
        ["Automated (Cron)", "Daily at 00:00 UTC", "ENABLE_CRON=true, CRON_TIMEZONE=UTC, schedule: '0 0 * * *'", "Runs all 6 scrapers sequentially, then chunks, embeds, and indexes"],
        ["Manual (Admin UI)", "On-demand", "POST /api/admin/scrape/trigger (JWT required)", "Admin clicks 'Run Full Scrape & Reindex' in dashboard. Status polled via GET /api/admin/scrape/status every 5s"],
        ["Manual (CLI)", "On-demand", "node src/pipeline.js", "Direct execution for debugging or emergency re-index"],
    ],
    col_widths=[2.5, 3, 5, 6]
)

add_sub_sub_heading("Failure Handling")
add_bullet("If any scraper throws an error, the pipeline logs the error and continues with remaining scrapers. Partial data is still processed.", bold_prefix="Graceful Degradation:")
add_bullet("If Puppeteer cannot load a page (timeout: 60s per page, navigation timeout), it retries the page load. If persistent, that data source is skipped.", bold_prefix="Timeout Handling:")
add_bullet("Pipeline status is tracked as a state machine: idle → running → success/error. The last status and error message are available via the admin API.", bold_prefix="Status Tracking:")
add_bullet("Admin dashboard shows lastScrapeAt timestamp. If the pipeline hasn't run in > 24 hours, monitoring alerts should fire.", bold_prefix="Staleness Detection:")

add_sub_heading("7.2  Re-indexing Strategy & Operational Triggers")

add_sub_sub_heading("Current Strategy: Full Re-index")
add_body(
    "Every pipeline run performs a complete re-index: deleteByQuery({ match_all: {} }) clears ALL documents "
    "from the Elasticsearch index, then re-embeds and re-indexes all freshly scraped chunks. This ensures:"
)
add_bullet("Stale job listings that have been removed from the UNIDO website are automatically purged", color=GREEN)
add_bullet("Updated job details, new postings, and content changes are immediately reflected", color=GREEN)
add_bullet("No index fragmentation or orphaned documents accumulate over time", color=GREEN)
add_bullet("The index is always a perfect mirror of the current UNIDO careers website", color=GREEN)

add_sub_sub_heading("Operational Triggers for Re-indexing")
add_styled_table(
    ["Trigger", "Method", "When to Use"],
    [
        ["Scheduled daily cron", "Automatic", "Standard operation — ensures data freshness within 24 hours"],
        ["Admin UI button", "Manual (admin dashboard)", "After UNIDO publishes urgent job postings or major content changes"],
        ["CLI pipeline run", "Manual (SSH/terminal)", "After infrastructure changes, index schema updates, or debugging"],
        ["Post-deployment", "CI/CD pipeline step", "After deploying code changes that affect chunking logic or index mapping"],
        ["Index corruption recovery", "Manual", "If Elasticsearch reports index errors or search returns unexpected results"],
    ],
    col_widths=[3.5, 3, 10]
)

add_sub_heading("7.3  Manual vs. Automated Maintenance Tasks")

add_styled_table(
    ["Task", "Type", "Frequency", "Owner", "Procedure"],
    [
        ["Data pipeline execution", "Automated", "Daily (midnight UTC)", "System (node-cron)", "No action needed. Monitor lastScrapeAt and pipeline status via admin dashboard"],
        ["Elasticsearch re-indexing", "Automated", "Part of pipeline", "System", "Included in daily pipeline. Full delete + re-index on every run"],
        ["MongoDB backups", "Automated", "Continuous (Atlas)", "MongoDB Atlas", "Point-in-time recovery available. Verify via Atlas console monthly"],
        ["SSL/TLS certificate renewal", "Automated", "90 days (App Service managed)", "Azure", "App Service manages Let's Encrypt certificates. Verify expiry quarterly"],
        ["Dependency security patches", "Manual", "Monthly", "Dev team", "Run npm audit, review CVEs, update package.json, test, deploy"],
        ["Node.js runtime update", "Manual", "Quarterly", "Dev team", "Update App Service runtime version when new Node.js LTS is released"],
        ["Monitor scraping accuracy", "Manual", "Monthly", "Operations", "Compare chatbot answers against live UNIDO website. Review blocked query logs for false positives"],
        ["Admin password rotation", "Manual", "90 days", "UNIDO admin", "Change password via admin API or database. JWT_SECRET rotation requires coordinated deploy"],
        ["Review blocked queries", "Manual", "Weekly", "Operations", "Export chat logs with status='fallback', review guardReason for false positives, tune guard patterns if needed"],
        ["Cost review", "Manual", "Monthly", "Operations", "Review Azure Cost Management dashboard. Check OpenAI token usage trends. Right-size resources"],
        ["Elasticsearch cluster health", "Manual", "Monthly", "Operations", "Check cluster status (green/yellow/red), index size, query latency via Elastic Cloud console"],
        ["Log review & cleanup", "Semi-automated", "Monthly", "Operations", "Review Azure Monitor logs. Ensure retention policies are active. Archive if needed"],
    ],
    col_widths=[3, 2, 2, 2.5, 7]
)

add_callout(
    "Operational overhead is intentionally low by design. The chatbot's architecture prioritises automation: "
    "daily data refresh, automated backups, auto-scaling, and built-in monitoring. Routine manual tasks are limited "
    "to monthly security reviews, dependency updates, and operational spot-checks. No database administration, "
    "index management, or content curation is required from the operations team.",
    "0066B3", "EBF5FF"
)


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
run = p.add_run("End of Document")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = GREY_DARK

add_horizontal_line("0066B3")

doc.add_paragraph()
add_body(
    "This document provides comprehensive, implementation-grounded answers to all client queries. "
    "All specifications, configurations, and cost estimates are derived directly from the application's "
    "codebase, Azure service documentation, and current pricing models (April 2026)."
)

doc.add_paragraph()
add_body(
    "MouriTech remains available for follow-up discussions, detailed architecture reviews, or hands-on "
    "demonstrations of any component described in this document."
)

doc.add_paragraph()
doc.add_paragraph()

# Readiness callout
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tbl)
cell = tbl.cell(0, 0)
cell.text = ""
set_cell_shading(cell, "0D2B52")
p = cell.paragraphs[0]
run = p.add_run("  Next Steps")
run.font.bold = True; run.font.size = Pt(13); run.font.color.rgb = BLUE_MID
items = [
    "UNIDO to review and approve data retention policies (Section 6)",
    "UNIDO to provision Azure resources per infrastructure requirements (Section 1)",
    "MouriTech to set up CI/CD pipelines once repository access is confirmed (Section 3)",
    "Joint session to finalise environment-specific configurations and credentials",
]
for item in items:
    p2 = cell.add_paragraph()
    r = p2.add_run(f"  →  {item}")
    r.font.size = Pt(10); r.font.color.rgb = WHITE

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by MouriTech Global Enterprise Solutions  |  Version 1.0  |  April 2026  |  Confidential")
run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = GREY_MID


# ── SAVE ──
output_path = os.path.join(os.path.dirname(__file__), "..", "docs", "UNIDO_Client_Queries_Response.docx")
doc.save(output_path)
print(f"Word document saved to: {os.path.abspath(output_path)}")
